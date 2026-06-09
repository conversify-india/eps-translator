#!/usr/bin/env python3
import sys
import os
import json
import re
import docx
from docx.oxml.ns import qn, nsmap
from docx.shared import Pt
from docx.oxml import parse_xml

# Register drawing and shape namespaces
nsmap['v'] = 'urn:schemas-microsoft-com:vml'
nsmap['wps'] = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
nsmap['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'
nsmap['wp'] = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

def normalize_text(text):
    if not text:
        return ""
    # Collapse multiple whitespaces and newlines into a single space
    return re.sub(r'\s+', ' ', text).strip()

def apply_run_font_scaling(run_obj, paragraph_obj, scale_factor):
    if scale_factor >= 1.0:
        return
    current_size = None
    if run_obj.font.size:
        current_size = run_obj.font.size.pt
    elif paragraph_obj.style and paragraph_obj.style.font and paragraph_obj.style.font.size:
        current_size = paragraph_obj.style.font.size.pt
    else:
        current_size = 11.0
    run_obj.font.size = Pt(current_size * scale_factor)

def replace_text_in_paragraph_single(p, key, replacement):
    # Safe check if key is in the text
    if key not in p.text:
        return False

    # Calculate length ratio and scale factor
    len_orig = len(key)
    len_trans = len(replacement)
    scale_factor = 1.0
    if len_orig > 5 and len_trans > len_orig:
        scale_factor = max(0.70, len_orig / len_trans)  # Down to 70% maximum shrinkage

    # Helper to apply font scaling to a run
    def scale_run(run_obj):
        apply_run_font_scaling(run_obj, p, scale_factor)

    # Exact paragraph match (modulo outer whitespace)
    if p.text.strip() == key.strip():
        if p.runs:
            for r in p.runs:
                scale_run(r)
            p.runs[0].text = replacement
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = replacement
        return True

    # Fallback if no runs exist
    if not p.runs:
        p.text = p.text.replace(key, replacement)
        return True

    # Run-by-run exact substring replacement
    for r in p.runs:
        if key in r.text:
            scale_run(r)
            r.text = r.text.replace(key, replacement)
            return True

    # Key is split across multiple runs
    char_map = []
    for r_idx, r in enumerate(p.runs):
        if r.text:
            for c_idx in range(len(r.text)):
                char_map.append((r_idx, c_idx))

    # Safety bounds check
    if len(char_map) != len(p.text):
        # Fallback to direct replacement if character mapping mismatch
        p.text = p.text.replace(key, replacement)
        return True

    start_idx = p.text.find(key)
    if start_idx == -1:
        return False

    end_idx = start_idx + len(key)
    
    involved_runs = sorted(list(set(char_map[i][0] for i in range(start_idx, end_idx))))
    if not involved_runs:
        # Fallback
        p.text = p.text.replace(key, replacement)
        return True

    first_r_idx, first_c_idx = char_map[start_idx]
    last_r_idx, last_c_idx = char_map[end_idx - 1]

    first_run = p.runs[first_r_idx]
    last_run = p.runs[last_r_idx]

    prefix = first_run.text[:first_c_idx]
    suffix = last_run.text[last_c_idx + 1:]

    # Scale font sizes in all involved runs
    for r_idx in involved_runs:
        scale_run(p.runs[r_idx])

    # Apply replacement
    first_run.text = prefix + replacement

    # Clear text in remaining runs involved in this match
    for r_idx in involved_runs[1:]:
        p.runs[r_idx].text = ""

    if first_r_idx != last_r_idx:
        last_run.text = suffix
    else:
        first_run.text += suffix

    return True

def replace_text_in_paragraph(p, key, replacement):
    limit = 10
    replaced = False
    while key in p.text and limit > 0:
        if replace_text_in_paragraph_single(p, key, replacement):
            replaced = True
        limit -= 1
    return replaced

def iter_block_paragraphs(parent):
    # Paragraphs in body or block parent (like table cell)
    if hasattr(parent, 'paragraphs'):
        for p in parent.paragraphs:
            yield p
    # Tables inside parent (recursive support for nested cells)
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_paragraphs(cell)

def adjust_table_row_heights(parent):
    # Process tables in the current block parent (like Body, Header, Cell) recursively
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            for row in table.rows:
                trPr = row._tr.get_or_add_trPr()
                trHeight = trPr.first_child_found_in("w:trHeight")
                if trHeight is not None:
                    hRule = trHeight.get(qn('w:hRule'))
                    if hRule == 'exact':
                        trHeight.set(qn('w:hRule'), 'atLeast')
                for cell in row.cells:
                    adjust_table_row_heights(cell)

def main():
    autofit = "--autofit" in sys.argv

    if len(sys.argv) < 4:
        print("Usage: python docx_translation_replacer.py <input_docx_path> <output_docx_path> <translation_json_path> [--autofit]")
        sys.exit(1)

    input_docx_path = sys.argv[1]
    output_docx_path = sys.argv[2]
    translation_json_path = sys.argv[3]

    if not os.path.exists(input_docx_path):
        print(f"Error: Input DOCX file not found at {input_docx_path}")
        sys.exit(1)
    if not os.path.exists(translation_json_path):
        print(f"Error: Translation JSON file not found at {translation_json_path}")
        sys.exit(1)

    # 1. Open DOCX for translation post-processing
    print("Loading generated DOCX structures...")
    try:
        doc = docx.Document(input_docx_path)
    except Exception as e:
        print(f"Error loading generated DOCX: {e}")
        sys.exit(1)

    # Change all table row height rules from exact to atLeast to prevent text clipping
    print("Adjusting table row height rules to allow expansion...")
    adjust_table_row_heights(doc)
    for section in doc.sections:
        if section.header:
            adjust_table_row_heights(section.header)
        if section.footer:
            adjust_table_row_heights(section.footer)

    # Auto-fit tables and set cell widths to auto to prevent vertical text wrapping (only if local pdf2docx fallback was used)
    if autofit:
        print("Auto-fitting tables and clearing narrow cell widths (local fallback mode)...")
        for table in doc.tables:
            table.autofit = True
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcW = tcPr.first_child_found_in("w:tcW")
                    if tcW is not None:
                        tcW.set(qn('w:w'), '0')
                        tcW.set(qn('w:type'), 'auto')

    # 2. Read translation map
    try:
        with open(translation_json_path, 'r', encoding='utf-8') as f:
            translation_map = json.load(f)
    except Exception as e:
        print(f"Error reading translation map: {e}")
        sys.exit(1)

    # 3. Prepare normalized keys and sort them by length descending
    normalized_map = {}
    original_map = {}
    
    for original, translated in translation_map.items():
        if not original or not translated:
            continue
        orig_cleaned = original.strip()
        trans_cleaned = translated.strip()
        if not orig_cleaned:
            continue
            
        original_map[orig_cleaned] = trans_cleaned
        norm_orig = normalize_text(orig_cleaned)
        if norm_orig:
            normalized_map[norm_orig] = trans_cleaned

    sorted_norm_keys = sorted(normalized_map.keys(), key=len, reverse=True)
    sorted_orig_keys = sorted(original_map.keys(), key=len, reverse=True)

    print(f"Translation map loaded: {len(original_map)} strings.")

    # 4. Iterate through all paragraphs in body, tables, headers, footers
    paragraphs = []
    
    # Body paragraphs and paragraphs inside tables
    for p in iter_block_paragraphs(doc):
        paragraphs.append(p)
        
    # Header and footer paragraphs
    for section in doc.sections:
        if section.header:
            for p in iter_block_paragraphs(section.header):
                paragraphs.append(p)
        if section.footer:
            for p in iter_block_paragraphs(section.footer):
                paragraphs.append(p)

    # 4.1 Collect paragraphs inside floating text boxes (body, headers, footers)
    from docx.text.paragraph import Paragraph
    txbx_count = 0
    
    # Body text boxes
    for txbx in doc.element.xpath('//w:txbxContent'):
        for p_el in txbx.xpath('.//w:p', namespaces=nsmap):
            paragraphs.append(Paragraph(p_el, doc))
            txbx_count += 1
            
    # Header/Footer text boxes
    for section in doc.sections:
        if section.header:
            for txbx in section.header._element.xpath('//w:txbxContent'):
                for p_el in txbx.xpath('.//w:p', namespaces=nsmap):
                    paragraphs.append(Paragraph(p_el, section.header))
                    txbx_count += 1
        if section.footer:
            for txbx in section.footer._element.xpath('//w:txbxContent'):
                for p_el in txbx.xpath('.//w:p', namespaces=nsmap):
                    paragraphs.append(Paragraph(p_el, section.footer))
                    txbx_count += 1
                    
    print(f"Extracted {txbx_count} paragraphs from text boxes for translation.")

    # 4.2 Set all text boxes (VML and DrawingML) to solid white background to prevent overlaps
    vml_count = 0
    # Process VML shapes in body
    for shape in doc.element.xpath('//v:shape'):
        shape.set('filled', 't')
        shape.set('fillcolor', '#ffffff')
        style = shape.get('style', '')
        if style:
            style_parts = [p.strip() for p in style.split(';') if p.strip()]
            new_style_parts = []
            for part in style_parts:
                if part.lower().startswith('fill:') or part.lower().startswith('fill-opacity:'):
                    continue
                new_style_parts.append(part)
            shape.set('style', ';'.join(new_style_parts))
        vml_count += 1
        
    # Process VML shapes in headers/footers
    for section in doc.sections:
        if section.header:
            for shape in section.header._element.xpath('//v:shape'):
                shape.set('filled', 't')
                shape.set('fillcolor', '#ffffff')
                style = shape.get('style', '')
                if style:
                    style_parts = [p.strip() for p in style.split(';') if p.strip()]
                    new_style_parts = []
                    for part in style_parts:
                        if part.lower().startswith('fill:') or part.lower().startswith('fill-opacity:'):
                            continue
                        new_style_parts.append(part)
                    shape.set('style', ';'.join(new_style_parts))
                vml_count += 1
        if section.footer:
            for shape in section.footer._element.xpath('//v:shape'):
                shape.set('filled', 't')
                shape.set('fillcolor', '#ffffff')
                style = shape.get('style', '')
                if style:
                    style_parts = [p.strip() for p in style.split(';') if p.strip()]
                    new_style_parts = []
                    for part in style_parts:
                        if part.lower().startswith('fill:') or part.lower().startswith('fill-opacity:'):
                            continue
                        new_style_parts.append(part)
                    shape.set('style', ';'.join(new_style_parts))
                vml_count += 1

    dml_count = 0
    # Process DrawingML shapes in body
    for spPr in doc.element.xpath('//wps:wsp/wps:spPr'):
        noFill = spPr.find(qn('a:noFill'))
        if noFill is not None:
            spPr.remove(noFill)
        solidFill = spPr.find(qn('a:solidFill'))
        if solidFill is not None:
            spPr.remove(solidFill)
        new_fill = parse_xml(
            '<a:solidFill xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
            '<a:srgbClr val="FFFFFF"/>'
            '</a:solidFill>'
        )
        spPr.append(new_fill)
        dml_count += 1

    # Process DrawingML shapes in headers/footers
    for section in doc.sections:
        if section.header:
            for spPr in section.header._element.xpath('//wps:wsp/wps:spPr'):
                noFill = spPr.find(qn('a:noFill'))
                if noFill is not None:
                    spPr.remove(noFill)
                solidFill = spPr.find(qn('a:solidFill'))
                if solidFill is not None:
                    spPr.remove(solidFill)
                new_fill = parse_xml(
                    '<a:solidFill xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                    '<a:srgbClr val="FFFFFF"/>'
                    '</a:solidFill>'
                )
                spPr.append(new_fill)
                dml_count += 1
        if section.footer:
            for spPr in section.footer._element.xpath('//wps:wsp/wps:spPr'):
                noFill = spPr.find(qn('a:noFill'))
                if noFill is not None:
                    spPr.remove(noFill)
                solidFill = spPr.find(qn('a:solidFill'))
                if solidFill is not None:
                    spPr.remove(solidFill)
                new_fill = parse_xml(
                    '<a:solidFill xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                    '<a:srgbClr val="FFFFFF"/>'
                    '</a:solidFill>'
                )
                spPr.append(new_fill)
                dml_count += 1

    print(f"Updated backgrounds to solid white: {vml_count} VML shapes, {dml_count} DrawingML shapes.")

    # 4.3 Remove full-page background images to prevent double-text overlays
    def remove_full_page_backgrounds(element):
        removed = 0
        drawings = element.xpath('.//w:drawing')
        for dwg in drawings:
            anchor = dwg.find(qn('wp:anchor'))
            if anchor is not None:
                behindDoc = anchor.get('behindDoc')
                extent = anchor.find(qn('wp:extent'))
                if extent is not None:
                    try:
                        cx = int(extent.get('cx'))
                        cy = int(extent.get('cy'))
                        width_in = cx / 914400
                        height_in = cy / 914400
                        # A4 is 8.27 x 11.69 inches. Letter is 8.5 x 11 inches.
                        # If behind text and large, it is a full-page background image
                        if (behindDoc in ('1', 'true')) and width_in > 6.0 and height_in > 9.0:
                            parent = dwg.getparent()
                            if parent is not None:
                                parent.remove(dwg)
                                removed += 1
                    except Exception:
                        pass
        return removed

    removed_bg_count = remove_full_page_backgrounds(doc.element)
    for section in doc.sections:
        if section.header:
            removed_bg_count += remove_full_page_backgrounds(section.header._element)
        if section.footer:
            removed_bg_count += remove_full_page_backgrounds(section.footer._element)
    print(f"Removed {removed_bg_count} full-page background images.")

    # 4.4 Apply white background shading to all text paragraphs to mask any other overlaps
    def apply_shading_to_paragraph(p):
        pPr = p._element.get_or_add_pPr()
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        new_shd = parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="FFFFFF"/>')
        pPr.append(new_shd)

    shaded_p_count = 0
    for p in paragraphs:
        if p.text.strip():
            apply_shading_to_paragraph(p)
            shaded_p_count += 1
    print(f"Applied white paragraph shading to {shaded_p_count} text paragraphs.")

    # 5. Apply search and replace
    print(f"Processing text replacement on {len(paragraphs)} paragraph blocks...")
    replaced_count = 0
    
    for p in paragraphs:
        if not p.text.strip():
            continue
            
        p_text_norm = normalize_text(p.text)
        replaced_p = False

        # Phase A: Exact paragraph match on normalized text
        for key in sorted_norm_keys:
            if p_text_norm == key:
                replacement = normalized_map[key]
                len_orig = len(key)
                len_trans = len(replacement)
                scale_factor = 1.0
                if len_orig > 5 and len_trans > len_orig:
                    scale_factor = max(0.70, len_orig / len_trans)

                if p.runs:
                    for r in p.runs:
                        apply_run_font_scaling(r, p, scale_factor)
                    p.runs[0].text = replacement
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = replacement
                replaced_p = True
                replaced_count += 1
                break

        if replaced_p:
            continue

        # Phase B: Substring match on original keys (un-normalized)
        for orig_key in sorted_orig_keys:
            if orig_key in p.text:
                replacement = original_map[orig_key]
                if replace_text_in_paragraph(p, orig_key, replacement):
                    replaced_p = True
                    replaced_count += 1

    print(f"Post-processing completed. Replaced {replaced_count} matches.")

    # Save output DOCX
    try:
        doc.save(output_docx_path)
        print(f"Successfully saved translated DOCX to: {output_docx_path}")
    except Exception as e:
        print(f"Error saving translated DOCX: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
