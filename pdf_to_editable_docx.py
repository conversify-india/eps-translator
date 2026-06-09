#!/usr/bin/env python3
"""
pdf_to_editable_docx.py — Exact-Replica PDF → Editable Word Converter

Strategy:
  1. Render each PDF page as a high-res background image
  2. Place the image as a full-page background in Word (via header)
  3. Extract all text spans with their exact (x, y) coordinates
  4. Overlay absolutely-positioned transparent text boxes on top
  5. Result: visually identical to original PDF, but text is editable

Usage:
    python pdf_to_editable_docx.py <input.pdf> <output.docx> [translations.json]
"""
import sys
import os
import io
import json
import re
import tempfile

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# ─── Constants ──────────────────────────────────────────
DPI = 150  # Background image render resolution
PTS_PER_INCH = 72
EMU_PER_PT = 12700
EMU_PER_INCH = 914400


def pts_to_emu(pts):
    return int(pts * EMU_PER_INCH / PTS_PER_INCH)


def _is_bold(flags, font_name):
    return bool(flags & (1 << 4)) or any(k in font_name.lower() for k in ["bold", "black", "heavy"])


def _is_italic(flags, font_name):
    return bool(flags & (1 << 1)) or any(k in font_name.lower() for k in ["italic", "oblique"])


def _clean_font_name(raw):
    """Clean PDF font name to a Word-safe name."""
    name = re.sub(r'^[A-Z]{6}\+', '', raw)
    name = re.sub(r'[,\-](Bold|Italic|Regular|Light|Medium|Black|Heavy|BoldItalic).*$', '', name, flags=re.IGNORECASE)
    font_map = {
        "TimesNewRoman": "Times New Roman", "CourierNew": "Courier New",
        "HelveticaNeue": "Arial", "Helvetica": "Arial", "ArialMT": "Arial",
    }
    for pdf_name, word_name in font_map.items():
        if pdf_name.lower() in name.lower():
            return word_name
    return name or "Arial"


def _color_tuple(c):
    if c is None:
        return (0, 0, 0)
    return ((c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF)


# ─── Extract text spans from a page ─────────────────────
def extract_spans(page, translation_map=None):
    """Extract all text spans with coordinates and formatting."""
    raw = page.get_text("dict")
    spans = []
    for block in raw.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                txt = span.get("text", "").strip()
                if not txt:
                    continue
                # Apply translation
                if translation_map:
                    if txt in translation_map:
                        txt = translation_map[txt]
                    else:
                        for orig, trans in translation_map.items():
                            if len(orig) > 3 and orig in txt:
                                txt = txt.replace(orig, trans)
                spans.append({
                    "text": txt,
                    "x": span["bbox"][0],
                    "y": span["bbox"][1],
                    "w": span["bbox"][2] - span["bbox"][0],
                    "h": span["bbox"][3] - span["bbox"][1],
                    "size": span.get("size", 10),
                    "font": _clean_font_name(span.get("font", "Arial")),
                    "bold": _is_bold(span.get("flags", 0), span.get("font", "")),
                    "italic": _is_italic(span.get("flags", 0), span.get("font", "")),
                    "color": _color_tuple(span.get("color", 0)),
                })
    return spans


# ─── Merge nearby spans into lines ──────────────────────
def merge_spans_to_lines(spans, y_threshold=3):
    """Group spans on the same visual line (within y_threshold pts)."""
    if not spans:
        return []
    sorted_spans = sorted(spans, key=lambda s: (round(s["y"] / y_threshold), s["x"]))
    lines = []
    current_line = [sorted_spans[0]]
    for s in sorted_spans[1:]:
        if abs(s["y"] - current_line[0]["y"]) < y_threshold:
            current_line.append(s)
        else:
            lines.append(current_line)
            current_line = [s]
    lines.append(current_line)
    return lines


# ─── Create a floating text box in Word XML ─────────────
def make_textbox_xml(text, x_emu, y_emu, w_emu, h_emu, font_name="Arial",
                     font_size_pt=10, bold=False, italic=False, color=(0, 0, 0)):
    """
    Create a Word VML floating text box at absolute position.
    The text box has no background/border, so the page image shows through.
    """
    bold_val = "on" if bold else "off"
    italic_val = "on" if italic else "off"
    hex_color = f"#{color[0]:02x}{color[1]:02x}{color[2]:02x}"
    font_size_hpt = int(font_size_pt * 2)  # half-points

    # Escape XML entities in text
    safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    xml_str = f'''
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:o="urn:schemas-microsoft-com:office:office"
         xmlns:w10="urn:schemas-microsoft-com:office:word">
      <w:rPr><w:noProof/></w:rPr>
      <w:pict>
        <v:shapetype id="_x0000_t202" coordsize="21600,21600" o:spt="202" path="m,l,21600r21600,l21600,xe">
          <v:stroke joinstyle="miter"/>
          <v:path gradientshapeok="t" o:connecttype="rect"/>
        </v:shapetype>
        <v:shape type="#_x0000_t202"
                 style="position:absolute;margin-left:{x_emu / EMU_PER_INCH:.4f}in;margin-top:{y_emu / EMU_PER_INCH:.4f}in;width:{w_emu / EMU_PER_INCH:.4f}in;height:{h_emu / EMU_PER_INCH:.4f}in;z-index:251;mso-position-horizontal-relative:page;mso-position-vertical-relative:page"
                 filled="f" stroked="f">
          <v:textbox style="mso-fit-shape-to-text:t" inset="0,0,0,0">
            <w:txbxContent>
              <w:p>
                <w:pPr>
                  <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>
                    <w:sz w:val="{font_size_hpt}"/>
                    <w:szCs w:val="{font_size_hpt}"/>
                    <w:b w:val="{bold_val}"/>
                    <w:i w:val="{italic_val}"/>
                    <w:color w:val="{hex_color[1:]}"/>
                  </w:rPr>
                  <w:t xml:space="preserve">{safe_text}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:shape>
      </w:pict>
    </w:r>'''
    return parse_xml(xml_str)


# ─── Set background image via header ────────────────────
def set_page_background_image(section, image_path, page_width_emu, page_height_emu):
    """
    Set a full-page background image by placing it in the header
    as an absolutely positioned image covering the entire page.
    """
    header = section.header
    header.is_linked_to_previous = False

    # Clear existing header content
    for p in header.paragraphs:
        p.clear()

    # Add image relationship
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    image_part_rId = header.part.relate_to(
        header.part._package.part_for(image_path) if hasattr(header.part._package, 'part_for') else None,
        RT.IMAGE
    ) if False else None

    # Use a simpler approach: add the image to the header paragraph
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    run = para.add_run()
    run.add_picture(image_path, width=Emu(page_width_emu), height=Emu(page_height_emu))

    # Position the image behind text using XML
    inline = run._element.findall(qn('w:drawing'))[0].find(qn('wp:inline'))
    if inline is not None:
        # Convert inline to anchor (behind text)
        anchor = _inline_to_anchor(inline, page_width_emu, page_height_emu)
        drawing = run._element.findall(qn('w:drawing'))[0]
        drawing.remove(inline)
        drawing.append(anchor)


def _inline_to_anchor(inline, cx, cy):
    """Convert a wp:inline element to a wp:anchor positioned at page origin, behind text."""
    nsmap = {
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    }

    # Get graphic element from inline
    graphic = inline.find(qn('a:graphic'))
    extent = inline.find(qn('wp:extent'))
    docPr = inline.find(qn('wp:docPr'))

    anchor_xml = f'''<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                    distT="0" distB="0" distL="0" distR="0"
                    simplePos="0" relativeHeight="0" behindDoc="1"
                    locked="1" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
    </wp:anchor>'''

    anchor = parse_xml(anchor_xml)
    if docPr is not None:
        anchor.append(docPr)
    else:
        anchor.append(parse_xml('<wp:docPr xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" id="1" name="Background"/>'))
    anchor.append(parse_xml('<wp:cNvGraphicFramePr xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"/>'))
    if graphic is not None:
        anchor.append(graphic)
    return anchor


# ─── Main conversion function ───────────────────────────
def convert_pdf_to_editable_docx(pdf_path, output_path, translation_map=None):
    """
    Convert PDF to editable Word document with exact visual layout.

    Each PDF page becomes one Word page with:
    - The rendered page image as background (behind text)
    - Editable text boxes at exact positions on top
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    fitz_doc = fitz.open(pdf_path)
    
    # Check if PDF is scanned or image-only
    scanned_pages = 0
    for page_idx in range(len(fitz_doc)):
        page = fitz_doc[page_idx]
        pw, ph = page.rect.width, page.rect.height
        text = page.get_text().strip()
        images = page.get_images(full=True)
        
        # 1. No text
        if not text:
            scanned_pages += 1
            continue
        # 2. Sparse text with images
        if len(text) < 150 and images:
            scanned_pages += 1
            continue
        # 3. Giant image covering > 85% of the page
        is_page_scanned = False
        for img_info in images:
            try:
                img_rects = page.get_image_bbox(img_info)
                rect = img_rects if isinstance(img_rects, fitz.Rect) else (img_rects[0] if img_rects else None)
                if rect:
                    img_w = rect.width
                    img_h = rect.height
                    if (img_w * img_h) > (pw * ph * 0.85):
                        if len(text) < 1200 or len(images) == 1:
                            is_page_scanned = True
                            break
            except:
                pass
        if is_page_scanned:
            scanned_pages += 1
            continue

    if scanned_pages > 0 or len(fitz_doc) == 0:
        print("Error: Scanned pages detected. This PDF requires OCR.")
        sys.exit(2)

    doc = Document()

    # Remove default paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    with tempfile.TemporaryDirectory() as tmpdir:
        for page_idx in range(len(fitz_doc)):
            page = fitz_doc[page_idx]
            pw_pts = page.rect.width
            ph_pts = page.rect.height
            pw_emu = pts_to_emu(pw_pts)
            ph_emu = pts_to_emu(ph_pts)

            print(f"  Page {page_idx + 1}/{len(fitz_doc)}: {pw_pts:.0f}×{ph_pts:.0f} pts")

            # ── 1. Extract text spans (from clean page before whiting out) ──
            spans = extract_spans(page, translation_map)
            lines = merge_spans_to_lines(spans)
            print(f"    Extracted: {len(spans)} spans → {len(lines)} lines")

            # ── 2. Draw solid white rectangles over original text spans to erase them from the background image ──
            for s in spans:
                # Add 1.5pt padding to completely cover anti-aliased font edges
                r = fitz.Rect(s["x"] - 1.5, s["y"] - 1.5, s["x"] + s["w"] + 1.5, s["y"] + s["h"] + 1.5)
                page.draw_rect(r, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)

            # ── 3. Render page as background image (now text-free) ──
            scale = DPI / PTS_PER_INCH
            mat = fitz.Matrix(scale, scale)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            bg_path = os.path.join(tmpdir, f"page_{page_idx}.png")
            pix.save(bg_path)
            print(f"    Rendered background: {pix.width}×{pix.height}px @ {DPI}dpi")

            # ── 3. Create Word section ──
            if page_idx == 0:
                section = doc.sections[0]
            else:
                section = doc.add_section()

            section.page_width = Emu(pw_emu)
            section.page_height = Emu(ph_emu)
            section.orientation = WD_ORIENT.LANDSCAPE if pw_pts > ph_pts else WD_ORIENT.PORTRAIT
            section.left_margin = Emu(0)
            section.right_margin = Emu(0)
            section.top_margin = Emu(0)
            section.bottom_margin = Emu(0)
            section.header_distance = Emu(0)
            section.footer_distance = Emu(0)

            # ── 4. Place background image in header ──
            try:
                set_page_background_image(section, bg_path, pw_emu, ph_emu)
                print(f"    Background image set")
            except Exception as e:
                print(f"    Warning: background image failed: {e}")
                # Fallback: insert as inline image
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(bg_path, width=Emu(pw_emu))

            # ── 5. Overlay text boxes ──
            # Add a single paragraph to hold all floating text boxes
            anchor_para = doc.add_paragraph()
            anchor_para.paragraph_format.space_before = Pt(0)
            anchor_para.paragraph_format.space_after = Pt(0)

            textbox_count = 0
            for line in lines:
                # Merge spans on same line into a single text box
                line_sorted = sorted(line, key=lambda s: s["x"])
                # Use first span for position, merge text
                full_text_parts = []
                for i, s in enumerate(line_sorted):
                    if i > 0:
                        gap = s["x"] - (line_sorted[i - 1]["x"] + line_sorted[i - 1]["w"])
                        if gap > 3:
                            full_text_parts.append("  ")
                    full_text_parts.append(s["text"])

                full_text = "".join(full_text_parts)
                if not full_text.strip():
                    continue

                # Position from first span
                first = line_sorted[0]
                last = line_sorted[-1]
                x_emu = pts_to_emu(first["x"])
                y_emu = pts_to_emu(first["y"])
                w_emu = pts_to_emu(last["x"] + last["w"] - first["x"] + 5)
                h_emu = pts_to_emu(max(s["h"] for s in line_sorted) + 2)

                try:
                    tb_element = make_textbox_xml(
                        full_text,
                        x_emu, y_emu, w_emu, h_emu,
                        font_name=first["font"],
                        font_size_pt=first["size"],
                        bold=first["bold"],
                        italic=first["italic"],
                        color=first["color"],
                    )
                    anchor_para._element.append(tb_element)
                    textbox_count += 1
                except Exception as e:
                    print(f"    Warning: textbox failed for '{full_text[:30]}': {e}")

            print(f"    Placed {textbox_count} text boxes")

    fitz_doc.close()
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")
    return output_path


# ─── CLI ────────────────────────────────────────────────
def main():
    if len(sys.argv) < 3:
        print("Usage: pdf_to_editable_docx.py <input.pdf> <output.docx> [translations.json]")
        sys.exit(1)

    pdf_path = sys.argv[1]
    output_path = sys.argv[2]
    tx_map = None
    if len(sys.argv) >= 4 and os.path.exists(sys.argv[3]):
        with open(sys.argv[3], "r", encoding="utf-8") as f:
            tx_map = json.load(f)

    convert_pdf_to_editable_docx(pdf_path, output_path, tx_map)


if __name__ == "__main__":
    main()
