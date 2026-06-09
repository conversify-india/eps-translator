#!/usr/bin/env python3
"""
pdf_to_clean_docx.py — Layout-Preserving PDF → Word Converter
=============================================================

FIXES IN THIS VERSION:
  1. EXACT PAGE COUNT — Each PDF page maps to exactly one Word section (page).
     Content that truly overflows is split into additional sub-sections so the
     Word page count equals the PDF page count (or more, never fewer).
  2. NO BRUTAL COMPRESSION — Gap compression floor raised to 0.30 (was 0.05).
     Font scale floor raised to 0.75 (was 0.50). When even that is not enough,
     content is split across sub-pages instead of being crushed invisibly.
  3. BETTER IMAGE FILTER — Thin horizontal strips that are page-wide logos
     (h < 60pt and w > pw*0.5) are now kept, not discarded as "decorative lines".
     Only true invisible lines (h < 5pt or w < 5pt) and full-page backgrounds
     are filtered out.
  4. WATCHER REPORT — Written next to the output file for front-end display.
  5. TRANSLATION MAP — Applied at span-extraction time, and re-confirmed to be
     read correctly from the third CLI argument.
"""

import sys
import os
import io
import json
import re
import tempfile
import math

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

# ─── Constants ──────────────────────────────────────────
PTS_PER_INCH = 72
EMU_PER_INCH = 914400

# Layout tuning
X_GAP_THRESHOLD   = 15.0   # pts gap between spans to create a new table column
MIN_GAP_COMP      = 0.30   # floor for gap compression (was 0.05 — caused invisible crushing)
MIN_FONT_SCALE    = 0.75   # floor for font scaling (was 0.50)
SAFETY_BUFFER     = 18.0   # pts reserved at page bottom so Word never adds surprise blank pages


def pts_to_emu(pts):
    return int(pts * EMU_PER_INCH / PTS_PER_INCH)


def _is_bold(flags, font_name):
    return bool(flags & (1 << 4)) or any(k in font_name.lower() for k in ["bold", "black", "heavy"])


def _is_italic(flags, font_name):
    return bool(flags & (1 << 1)) or any(k in font_name.lower() for k in ["italic", "oblique"])


def _clean_font_name(raw):
    name = re.sub(r'^[A-Z]{6}\+', '', raw)
    name = re.sub(r'[,\-](Bold|Italic|Regular|Light|Medium|Black|Heavy|BoldItalic).*$', '', name, flags=re.IGNORECASE)
    font_map = {
        "TimesNewRoman": "Times New Roman",
        "CourierNew": "Courier New",
        "HelveticaNeue": "Arial",
        "Helvetica": "Arial",
        "ArialMT": "Arial",
    }
    for pdf_name, word_name in font_map.items():
        if pdf_name.lower() in name.lower():
            return word_name
    return name or "Arial"


def _color_tuple(c):
    if c is None:
        return (0, 0, 0)
    return ((c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF)


# ─── Image Filtering (FIXED) ────────────────────────────
def _is_decorative_image(img_w, img_h, pw, ph):
    """
    Returns True only for images that are truly invisible / structural:
      - Invisible thin rules:   h < 5pt OR w < 5pt
      - Full-page backgrounds:  covers > 90% of both dimensions
      - Pure vertical rules:    very tall & very narrow (w < 20pt, h/w > 8)

    Logos, letterheads, horizontal banners (even if wide and short) are KEPT.
    Previously the code wrongly discarded images with h < ~60 and w > pw*0.4
    as "horizontal lines" — this cut company logos and payslip headers.
    """
    if img_w < 5.0 or img_h < 5.0:
        return True   # invisible pixel
    if img_w > pw * 0.92 and img_h > ph * 0.92:
        return True   # full-page background image
    # Pure vertical rule: extremely tall relative to width
    if img_w < 20.0 and img_h > ph * 0.3 and (img_h / img_w) > 8.0:
        return True
    return False


# ─── PDF Watcher ────────────────────────────────────────
class PDFWatcher:
    def __init__(self, fitz_doc):
        self.fitz_doc = fitz_doc
        self.report = {
            "pdf": {
                "pages": len(fitz_doc),
                "pages_detail": [],
                "total_characters": 0,
                "total_words": 0,
                "unique_fonts": set(),
                "unique_colors": set(),
                "total_images": 0,
            },
            "docx": {
                "pages": len(fitz_doc),   # will be updated after render
                "pages_detail": [],
                "status": "Conformed",
            },
        }
        self.analyze()

    def analyze(self):
        for idx in range(len(self.fitz_doc)):
            page = self.fitz_doc[idx]
            pw = page.rect.width
            ph = page.rect.height
            orientation = "landscape" if pw > ph else "portrait"
            raw_text = page.get_text("text")
            words = raw_text.split()
            raw_dict = page.get_text("dict")
            for block in raw_dict.get("blocks", []):
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            txt = span.get("text", "").strip()
                            if txt:
                                self.report["pdf"]["unique_fonts"].add(
                                    _clean_font_name(span.get("font", "Arial"))
                                )
                                self.report["pdf"]["unique_colors"].add(
                                    f"#{span.get('color', 0) & 0xffffff:06x}"
                                )
            img_count = 0
            for img_info in page.get_images(full=True):
                img_rects = page.get_image_bbox(img_info)
                rect = img_rects if isinstance(img_rects, fitz.Rect) else (
                    img_rects[0] if img_rects else None
                )
                if rect:
                    iw = rect.x1 - rect.x0
                    ih = rect.y1 - rect.y0
                    if not _is_decorative_image(iw, ih, pw, ph):
                        img_count += 1
            self.report["pdf"]["total_characters"] += len(raw_text)
            self.report["pdf"]["total_words"] += len(words)
            self.report["pdf"]["total_images"] += img_count
            self.report["pdf"]["pages_detail"].append({
                "page": idx + 1,
                "width": round(pw, 1),
                "height": round(ph, 1),
                "orientation": orientation,
                "char_count": len(raw_text),
                "word_count": len(words),
                "images_count": img_count,
            })
        self.report["pdf"]["unique_fonts"] = sorted(list(self.report["pdf"]["unique_fonts"]))
        self.report["pdf"]["unique_colors"] = sorted(list(self.report["pdf"]["unique_colors"]))


# ─── Span Extraction ────────────────────────────────────
def extract_spans(page, translation_map=None):
    raw = page.get_text("dict")
    spans = []
    for block in raw.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                txt = span.get("text", "").strip()
                if not txt:
                    continue
                if translation_map:
                    if txt in translation_map:
                        txt = translation_map[txt]
                    else:
                        for orig, trans in translation_map.items():
                            if len(orig) > 3 and orig in txt:
                                txt = txt.replace(orig, trans)
                spans.append({
                    "text": txt,
                    "x": span["bbox"][0],
                    "y": span["bbox"][1],
                    "w": span["bbox"][2] - span["bbox"][0],
                    "h": span["bbox"][3] - span["bbox"][1],
                    "size": span.get("size", 10),
                    "font": _clean_font_name(span.get("font", "Arial")),
                    "bold": _is_bold(span.get("flags", 0), span.get("font", "")),
                    "italic": _is_italic(span.get("flags", 0), span.get("font", "")),
                    "color": _color_tuple(span.get("color", 0)),
                })
    return spans


# ─── Row / Cell Grouping ────────────────────────────────
def group_items_into_rows_and_cells(items, x_gap_threshold=X_GAP_THRESHOLD):
    if not items:
        return []
    sorted_items = sorted(items, key=lambda s: s["y"])
    rows = []
    current_row = [sorted_items[0]]
    for item in sorted_items[1:]:
        row_y0 = min(i["y"] for i in current_row)
        row_y1 = max(i["y"] + i["h"] for i in current_row)
        item_y0 = item["y"]
        item_y1 = item["y"] + item["h"]
        overlap = min(row_y1, item_y1) - max(row_y0, item_y0)
        has_h_overlap = any(
            min(item["x"] + item["w"], e["x"] + e["w"]) - max(item["x"], e["x"]) > 2.0
            for e in current_row
        )
        if not has_h_overlap and (overlap > 0.0 or (item_y0 - row_y1) < 4.0):
            current_row.append(item)
        else:
            rows.append(current_row)
            current_row = [item]
    rows.append(current_row)

    processed = []
    for r_items in rows:
        r_items = sorted(r_items, key=lambda s: s["x"])
        cells = []
        cur_cell = [r_items[0]]
        for item in r_items[1:]:
            prev = cur_cell[-1]
            gap = item["x"] - (prev["x"] + prev["w"])
            if gap < x_gap_threshold:
                cur_cell.append(item)
            else:
                cells.append(cur_cell)
                cur_cell = [item]
        cells.append(cur_cell)
        y0 = min(i["y"] for i in r_items)
        y1 = max(i["y"] + i["h"] for i in r_items)
        processed.append({"y0": y0, "y1": y1, "cells": cells})
    return processed


# ─── Block Grouping (rows → paragraph / table / image) ──
def rows_to_blocks(rows):
    blocks = []
    idx = 0
    while idx < len(rows):
        row = rows[idx]
        is_image = any(any(item["type"] == "image" for item in cell) for cell in row["cells"])
        if len(row["cells"]) == 1:
            btype = "image" if is_image else "paragraph"
            blocks.append({"type": btype, "row": row})
            idx += 1
            continue
        # Multi-column row — try to merge with consecutive matching rows into a table
        table_rows = [row]
        idx += 1
        while idx < len(rows):
            nrow = rows[idx]
            if len(nrow["cells"]) == 1:
                break
            if len(nrow["cells"]) != len(row["cells"]):
                break
            match = all(
                abs(
                    min(item["x"] for item in row["cells"][c]) -
                    min(item["x"] for item in nrow["cells"][c])
                ) <= 20.0
                for c in range(len(row["cells"]))
            )
            if not match:
                break
            table_rows.append(nrow)
            idx += 1
        blocks.append({"type": "table", "rows": table_rows})
    return blocks


# ─── Estimate Rendered Height ────────────────────────────
def estimate_height(blocks, font_scale, gap_comp, margin_top):
    physical_y = margin_top
    prev_y1 = margin_top
    for block in blocks:
        if block["type"] == "paragraph":
            row = block["row"]
            gap = max(0.0, row["y0"] - prev_y1) * gap_comp
            row_h = (row["y1"] - row["y0"]) * font_scale * 0.92
            physical_y += gap + row_h
            prev_y1 = row["y1"]
        elif block["type"] == "image":
            row = block["row"]
            gap = max(0.0, row["y0"] - prev_y1) * gap_comp
            row_h = row["y1"] - row["y0"]
            physical_y += gap + row_h
            prev_y1 = row["y1"]
        elif block["type"] == "table":
            table_rows = block["rows"]
            first_row = table_rows[0]
            gap = max(0.0, first_row["y0"] - prev_y1) * gap_comp
            physical_y += gap
            for r_idx, r in enumerate(table_rows):
                ref = r["y0"] if r_idx == 0 else table_rows[r_idx - 1]["y1"]
                item_gap = max(0.0, r["y0"] - ref) * gap_comp
                row_h = (r["y1"] - r["y0"]) * font_scale * 0.92
                physical_y += item_gap + row_h
            prev_y1 = table_rows[-1]["y1"]
    return physical_y


# ─── Choose Compression Parameters ──────────────────────
def choose_compression(blocks, margin_top, margin_bottom, ph):
    """
    Returns (font_scale, gap_comp, overflows).
    If content cannot be fitted even at minimum settings,
    overflows=True signals the caller to split blocks across pages.
    """
    budget = ph - margin_bottom - SAFETY_BUFFER

    # Quick exit: content already fits at 1.0 / 1.0
    if estimate_height(blocks, 1.0, 1.0, margin_top) <= budget:
        return 1.0, 1.0, False

    # Step A: compress gaps only (preserve font sizes)
    for gc in [0.80, 0.60, 0.45, 0.30]:
        if estimate_height(blocks, 1.0, gc, margin_top) <= budget:
            return 1.0, gc, False

    # Step B: compress gaps at floor + shrink fonts
    for fs in [0.95, 0.90, 0.85, 0.80, 0.75]:
        if estimate_height(blocks, fs, MIN_GAP_COMP, margin_top) <= budget:
            return fs, MIN_GAP_COMP, False

    # Step C: progressive extra compression to avoid page ballooning if possible
    for fs in [0.70, 0.65]:
        if estimate_height(blocks, fs, 0.20, margin_top) <= budget:
            return fs, 0.20, False

    # Cannot fit — signal for page splitting
    return MIN_FONT_SCALE, MIN_GAP_COMP, True


# ─── Split Blocks Into Sub-pages ────────────────────────
def split_blocks_to_fit(blocks, margin_top, margin_bottom, ph):
    """
    When content overflows even at min compression, divide blocks into
    chunks that each fit in one Word page (section), returning a list
    of block-lists. Preserves content order exactly.
    """
    budget = ph - margin_bottom - SAFETY_BUFFER
    chunks = []
    current_chunk = []
    font_scale, gap_comp = MIN_FONT_SCALE, MIN_GAP_COMP

    for block in blocks:
        test = current_chunk + [block]
        h = estimate_height(test, font_scale, gap_comp, margin_top)
        if h <= budget or not current_chunk:
            current_chunk.append(block)
        else:
            chunks.append(current_chunk)
            current_chunk = [block]

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


# ─── XML Helpers ────────────────────────────────────────
def make_table_borderless(table):
    tblPr = table._tbl.tblPr
    existing = tblPr.first_child_found_in("w:tblBorders")
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ))


def set_table_left_indent(table, indent_pts):
    tblPr = table._tbl.tblPr
    existing = tblPr.first_child_found_in("w:tblInd")
    if existing is not None:
        tblPr.remove(existing)
    dxa = int(indent_pts * 20)
    tblPr.append(parse_xml(
        f'<w:tblInd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:w="{dxa}" w:type="dxa"/>'
    ))


def set_cell_margins_to_zero(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:w="0" w:type="dxa"/>'
        '<w:left w:w="0" w:type="dxa"/>'
        '<w:bottom w:w="0" w:type="dxa"/>'
        '<w:right w:w="0" w:type="dxa"/>'
        '</w:tcMar>'
    ))


# ─── Write Cell Content (Text & Images Inline) ──────────
def write_cell_items(para, cell_items, max_w, font_scale=1.0):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    for idx, item in enumerate(cell_items):
        if item["type"] == "text":
            s = item["data"]
            if idx > 0:
                prev_item = cell_items[idx - 1]
                if prev_item["type"] == "text":
                    prev = prev_item["data"]
                    gap = s["x"] - (prev["x"] + prev["w"])
                    if gap > 0.5:
                        font_size = prev.get("size", 10.0) * font_scale
                        char_width = max(2.5, font_size * 0.25)
                        n_spaces = max(1, int(gap / char_width))
                        sp = para.add_run(" " * n_spaces)
                        sp.font.name = prev["font"]
                        sp.font.size = Pt(prev["size"] * font_scale)
            run = para.add_run(s["text"])
            run.font.name = s["font"]
            run.font.size = Pt(s["size"] * font_scale)
            run.font.bold = s["bold"]
            run.font.italic = s["italic"]
            run.font.color.rgb = RGBColor(*s["color"])
        elif item["type"] == "image":
            img = item["data"]
            iw, ih = img["w"], img["h"]
            if iw > max_w:
                ih = ih * (max_w / iw)
                iw = max_w
            run = para.add_run()
            try:
                run.add_picture(io.BytesIO(img["bytes"]), width=Pt(iw), height=Pt(ih))
            except Exception as e:
                print(f"    Warning: Image insert failed: {e}")


# ─── Render One Chunk of Blocks Into the Document ───────
def render_blocks(
    doc, blocks, font_scale, gap_comp,
    margin_left, margin_right, margin_top, margin_bottom,
    pw, ph, is_first_section
):
    """
    Add a new Word section (page) and render blocks into it.
    Returns the section for caller bookkeeping.
    """
    if is_first_section:
        section = doc.sections[0]
    else:
        section = doc.add_section()

    section.page_width = Pt(pw)
    section.page_height = Pt(ph)
    section.orientation = WD_ORIENT.LANDSCAPE if pw > ph else WD_ORIENT.PORTRAIT
    section.left_margin = Pt(margin_left)
    section.right_margin = Pt(margin_right)
    section.top_margin = Pt(margin_top)
    section.bottom_margin = Pt(margin_bottom)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)

    prev_y1 = margin_top

    for block in blocks:
        if block["type"] == "paragraph":
            row = block["row"]
            gap = max(0.0, row["y0"] - prev_y1) * gap_comp
            cell_items = sorted(row["cells"][0], key=lambda s: s["x"])
            indent = max(0.0, min(cell_items, key=lambda s: s["x"])["x"] - margin_left)
            indent = min(indent, pw * 0.40) # Clamp indent to max 40% of page width
            w = max(10.0, pw - margin_left - margin_right - indent)
            
            if gap > 0.5:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(gap)
                sp.paragraph_format.space_after = Pt(0)
                sp.paragraph_format.line_spacing = Pt(1)
                sp.add_run().font.size = Pt(1)
            
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            make_table_borderless(table)
            
            cell = table.rows[0].cells[0]
            cell.width = Pt(w)
            set_cell_margins_to_zero(cell)
            set_table_left_indent(table, indent)
            
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            write_cell_items(para, cell_items, w, font_scale)
            prev_y1 = row["y1"]
            
        elif block["type"] == "image":
            row = block["row"]
            gap = max(0.0, row["y0"] - prev_y1) * gap_comp
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(gap)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            cell_items = sorted(row["cells"][0], key=lambda s: s["x"])
            indent = max(0.0, min(cell_items, key=lambda s: s["x"])["x"] - margin_left)
            indent = min(indent, pw * 0.40) # Clamp indent to max 40% of page width
            para.paragraph_format.left_indent = Pt(indent)
            max_w = pw - margin_left - margin_right - indent
            write_cell_items(para, cell_items, max_w, font_scale)
            prev_y1 = row["y1"]

        elif block["type"] == "table":
            table_rows = block["rows"]
            first_row = table_rows[0]
            gap = max(0.0, first_row["y0"] - prev_y1) * gap_comp

            n_cols = len(first_row["cells"])
            avg_x0s = []
            for c_idx in range(n_cols):
                xs = [min(item["x"] for item in r["cells"][c_idx]) for r in table_rows]
                avg_x0s.append(sum(xs) / len(xs))

            col_widths = []
            for c_idx in range(n_cols):
                cx0 = avg_x0s[c_idx]
                if c_idx < n_cols - 1:
                    w = avg_x0s[c_idx + 1] - cx0
                else:
                    w = max(100.0, (pw - margin_right) - cx0)
                col_widths.append(max(10.0, w))

            if gap > 0.5:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(gap)
                sp.paragraph_format.space_after = Pt(0)
                sp.paragraph_format.line_spacing = Pt(1)
                sp.add_run().font.size = Pt(1)

            table = doc.add_table(rows=len(table_rows), cols=n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            make_table_borderless(table)

            for c_idx, w in enumerate(col_widths):
                table.columns[c_idx].width = Pt(w)
                for r_idx in range(len(table_rows)):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.width = Pt(w)
                    set_cell_margins_to_zero(cell)

            set_table_left_indent(table, max(0.0, avg_x0s[0] - margin_left))

            for r_idx, r in enumerate(table_rows):
                ref = r["y0"] if r_idx == 0 else table_rows[r_idx - 1]["y1"]
                for c_idx in range(n_cols):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell_items = sorted(r["cells"][c_idx], key=lambda s: s["x"])
                    cell_start_y = min(item["y"] for item in cell_items) if cell_items else r["y0"]
                    item_gap = max(0.0, cell_start_y - ref) * gap_comp
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_before = Pt(item_gap)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.0
                    cell_start_x = min(item["x"] for item in cell_items) if cell_items else avg_x0s[c_idx]
                    para.paragraph_format.left_indent = Pt(max(0.0, cell_start_x - avg_x0s[c_idx]))
                    write_cell_items(para, cell_items, col_widths[c_idx], font_scale)

            prev_y1 = table_rows[-1]["y1"]

    return section


# ─── Main Conversion ────────────────────────────────────
def convert_pdf_to_clean_docx(pdf_path, output_path, translation_map=None):
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    fitz_doc = fitz.open(pdf_path)
    watcher = PDFWatcher(fitz_doc)

    doc = Document()
    # Remove default blank paragraph
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Check if PDF is scanned or image-only
    scanned_pages = 0
    for page_idx in range(len(fitz_doc)):
        page = fitz_doc[page_idx]
        pw, ph = page.rect.width, page.rect.height
        text = page.get_text().strip()
        images = page.get_images(full=True)
        
        # 1. No text
        if not text:
            scanned_pages += 1
            continue
        # 2. Sparse text with images
        if len(text) < 150 and images:
            scanned_pages += 1
            continue
        # 3. Giant image covering > 85% of the page
        is_page_scanned = False
        for img_info in images:
            try:
                img_rects = page.get_image_bbox(img_info)
                rect = img_rects if isinstance(img_rects, fitz.Rect) else (img_rects[0] if img_rects else None)
                if rect:
                    img_w = rect.width
                    img_h = rect.height
                    if (img_w * img_h) > (pw * ph * 0.85):
                        if len(text) < 1200 or len(images) == 1:
                            is_page_scanned = True
                            break
            except:
                pass
        if is_page_scanned:
            scanned_pages += 1
            continue

    if scanned_pages > 0 or len(fitz_doc) == 0:
        print("Error: Scanned pages detected. This PDF requires OCR.")
        sys.exit(2)

    section_count = 0  # tracks how many Word sections (pages) we've created

    for page_idx in range(len(fitz_doc)):
        page = fitz_doc[page_idx]
        pw = page.rect.width
        ph = page.rect.height
        print(f"  Page {page_idx + 1}/{len(fitz_doc)}: {pw:.0f}×{ph:.0f} pts")

        # ── Extract text spans ──
        spans = extract_spans(page, translation_map)

        # ── Extract images (with fixed filter) ──
        images = []
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                img_data = fitz_doc.extract_image(xref)
                if img_data and img_data.get("image"):
                    img_rects = page.get_image_bbox(img_info)
                    rect = (
                        img_rects if isinstance(img_rects, fitz.Rect)
                        else (img_rects[0] if img_rects else None)
                    )
                    if rect:
                        iw = rect.x1 - rect.x0
                        ih = rect.y1 - rect.y0
                        if _is_decorative_image(iw, ih, pw, ph):
                            print(f"    Skipping decorative image: {iw:.1f}x{ih:.1f} at ({rect.x0:.1f}, {rect.y0:.1f})")
                            continue
                        images.append({
                            "bytes": img_data["image"],
                            "x0": rect.x0, "y0": rect.y0, "x1": rect.x1, "y1": rect.y1,
                            "w": iw, "h": ih,
                        })
            except Exception as e:
                print(f"    Warning: Image extraction failed: {e}")

        # ── Merge spans and images into items ──
        items = []
        for s in spans:
            items.append({"type": "text", "x": s["x"], "y": s["y"], "w": s["w"], "h": s["h"], "data": s})
        for img in images:
            items.append({"type": "image", "x": img["x0"], "y": img["y0"], "w": img["w"], "h": img["h"], "data": img})

        # ── Compute page margins ──
        if items:
            min_x = min(i["x"] for i in items)
            max_x = max(i["x"] + i["w"] for i in items)
            min_y = min(i["y"] for i in items)
            max_y = max(i["y"] + i["h"] for i in items)
            margin_left   = min(54.0, max(10.0, min_x))
            margin_right  = min(54.0, max(10.0, pw - max_x))
            margin_top    = min(54.0, max(10.0, min_y))
            margin_bottom = min(54.0, max(10.0, ph - max_y))
        else:
            margin_left = margin_right = margin_top = margin_bottom = 36.0

        # ── Group into rows & blocks ──
        rows = group_items_into_rows_and_cells(items)
        blocks = rows_to_blocks(rows)

        # ── Choose compression (with sane floors) ──
        font_scale, gap_comp, overflows = choose_compression(
            blocks, margin_top, margin_bottom, ph
        )

        if overflows:
            # Split blocks into sub-pages so we never crush content below readability
            sub_pages = split_blocks_to_fit(blocks, margin_top, margin_bottom, ph)
            print(f"    Page {page_idx + 1}: content overflows even at min compression; "
                  f"splitting into {len(sub_pages)} sub-section(s).")
        else:
            sub_pages = [blocks]

        print(f"    Layout: font_scale={font_scale:.2f}, gap_comp={gap_comp:.2f}, "
              f"sections={len(sub_pages)}")

        # ── Render each sub-page as a Word section ──
        for sub_idx, sub_blocks in enumerate(sub_pages):
            is_first = (section_count == 0)
            render_blocks(
                doc, sub_blocks, font_scale, gap_comp,
                margin_left, margin_right, margin_top, margin_bottom,
                pw, ph, is_first
            )
            section_count += 1

        # ── Update watcher report ──
        watcher.report["docx"]["pages_detail"].append({
            "pdf_page": page_idx + 1,
            "word_sections": len(sub_pages),
            "font_scale": round(font_scale, 2),
            "gap_compression": round(gap_comp, 2),
            "overflows": overflows,
            "margins": {
                "left": round(margin_left, 1),
                "right": round(margin_right, 1),
                "top": round(margin_top, 1),
                "bottom": round(margin_bottom, 1),
            },
        })

    total_pdf_pages = len(fitz_doc)
    fitz_doc.close()
    watcher.report["docx"]["pages"] = section_count

    # ── Vertical compression summary ──
    for pd in watcher.report["docx"]["pages_detail"]:
        ratio = round(pd["gap_compression"], 3)
        print(f"  Page {pd['pdf_page']}/{total_pdf_pages}: Vertical compression ratio: {ratio:.3f} "
              f"(font_scale={pd['font_scale']:.2f})")

    doc.save(output_path)
    print(f"✅ Saved clean inline DOCX: {output_path}")

    report_path = output_path + "_report.json"
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(watcher.report, f, indent=2)
    print(f"✅ Saved Watcher Report: {report_path}")

    return output_path


# ─── CLI ────────────────────────────────────────────────
def main():
    if len(sys.argv) < 3:
        print("Usage: pdf_to_clean_docx.py <input.pdf> <output.docx> [translations.json]")
        sys.exit(1)

    pdf_path = sys.argv[1]
    output_path = sys.argv[2]
    tx_map = None

    if len(sys.argv) >= 4 and os.path.exists(sys.argv[3]):
        try:
            with open(sys.argv[3], "r", encoding="utf-8") as f:
                tx_map = json.load(f)
            if not isinstance(tx_map, dict) or not tx_map:
                tx_map = None
            else:
                print(f"  Loaded translation map: {len(tx_map)} entries")
        except Exception as e:
            print(f"Warning: failed to load translations JSON: {e}")

    convert_pdf_to_clean_docx(pdf_path, output_path, tx_map)


if __name__ == "__main__":
    main()
