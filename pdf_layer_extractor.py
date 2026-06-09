#!/usr/bin/env python3
"""
pdf_layer_extractor.py — Layer-by-Layer PDF to Editable Word Reconstructor

Extracts each layer of a PDF page-by-page using PyMuPDF (fitz):
  Layer 1: Page geometry (size, margins)
  Layer 2: Text blocks with full formatting (font, size, color, bold/italic)
  Layer 3: Embedded images
  Layer 4: Tables (detected from aligned text blocks)
  Layer 5: Background graphics rendered as raster if needed

Outputs a fully editable .docx file using python-docx.

Usage:
    python pdf_layer_extractor.py <input_pdf> <output_docx> [<translation_json>]
"""
import sys
import os
import io
import json
import base64
import tempfile
import re

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree

def make_table_borderless(table):
    tblPr = table._tbl.tblPr
    existing = tblPr.first_child_found_in("w:tblBorders")
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ))

def set_table_left_indent(table, indent_pts):
    tblPr = table._tbl.tblPr
    existing = tblPr.first_child_found_in("w:tblInd")
    if existing is not None:
        tblPr.remove(existing)
    dxa = int(indent_pts * 20)
    tblPr.append(parse_xml(
        f'<w:tblInd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:w="{dxa}" w:type="dxa"/>'
    ))

def set_cell_margins_to_zero(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:w="0" w:type="dxa"/>'
        '<w:left w:w="0" w:type="dxa"/>'
        '<w:bottom w:w="0" w:type="dxa"/>'
        '<w:right w:w="0" w:type="dxa"/>'
        '</w:tcMar>'
    ))

# ─────────────────────────────────────────────
# Utility helpers
# ─────────────────────────────────────────────

def pts_to_emu(pts):
    """Points (1/72 inch) → EMUs (914400 per inch)."""
    return int(pts * 914400 / 72)

def _color_tuple(color_int):
    """Convert fitz integer color to (R, G, B) 0-255 tuple."""
    if color_int is None:
        return (0, 0, 0)
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return (r, g, b)

def _is_bold(flags, font_name):
    flag_bold = bool(flags & (1 << 4))
    name_bold = any(kw in font_name.lower() for kw in ["bold", "black", "heavy", "demi"])
    return flag_bold or name_bold

def _is_italic(flags, font_name):
    flag_italic = bool(flags & (1 << 1))
    name_italic = any(kw in font_name.lower() for kw in ["italic", "oblique", "slant"])
    return flag_italic or name_italic

def _align_from_bbox(bbox, page_width):
    """Heuristically determine paragraph alignment from text position."""
    x0, _, x1, _ = bbox
    center = (x0 + x1) / 2
    page_center = page_width / 2
    left_margin = page_width * 0.15
    right_margin = page_width * 0.85
    if x0 > right_margin:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if abs(center - page_center) < page_width * 0.08 and x0 > left_margin:
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT

# ─────────────────────────────────────────────
# Table detection helpers
# ─────────────────────────────────────────────

def detect_tables(blocks, page_width, snap_x=10, snap_y=6):
    """
    Detect tabular text blocks by grouping into rows/columns.
    Returns: (table_groups, remaining_blocks)
      table_groups: list of { rows: [[cell_block, ...], ...], col_positions: [...], y_min: ..., y_max: ... }
      remaining_blocks: blocks that are NOT part of a table
    """
    if not blocks:
        return [], []

    # Group blocks into rows by Y proximity
    rows_map = {}
    for b in blocks:
        y_key = round(b["y0"] / snap_y) * snap_y
        rows_map.setdefault(y_key, []).append(b)

    sorted_rows = sorted(rows_map.items())

    # Identify which row keys are table rows
    table_row_ys = set()
    for y_key, row_blocks in sorted_rows:
        if len(row_blocks) >= 2:
            row_blocks_sorted = sorted(row_blocks, key=lambda b: b["x0"])
            xs = [b["x0"] for b in row_blocks_sorted]
            if max(xs) - min(xs) > page_width * 0.12: # slightly lower threshold for table detection
                table_row_ys.add(y_key)

    # Group consecutive table rows into separate tables, splitting on paragraphs or large gaps
    table_groups = []
    current_table_rows = []
    non_table_blocks = []

    for y_key, row_blocks in sorted_rows:
        if y_key in table_row_ys:
            # Check if there is a large gap from the previous row in the current table
            if current_table_rows:
                prev_y = current_table_rows[-1][0]
                if y_key - prev_y > 45.0: # vertical gap > 45pt
                    table_groups.append(current_table_rows)
                    current_table_rows = []
            current_table_rows.append((y_key, row_blocks))
        else:
            if current_table_rows:
                table_groups.append(current_table_rows)
                current_table_rows = []
            non_table_blocks.extend(row_blocks)

    if current_table_rows:
        table_groups.append(current_table_rows)

    # Convert each group of table rows into a table matrix
    final_table_groups = []
    for group in table_groups:
        group_blocks = []
        for y_key, r_blocks in group:
            group_blocks.extend(r_blocks)

        if not group_blocks:
            continue

        # Determine column positions for this specific table group
        all_x0s = sorted(set(round(b["x0"] / snap_x) * snap_x for b in group_blocks))
        
        row_data = {}
        for b in group_blocks:
            y_key = round(b["y0"] / snap_y) * snap_y
            x_key = min(all_x0s, key=lambda cx: abs(cx - b["x0"]))
            row_data.setdefault(y_key, {})[x_key] = b

        sorted_group_ys = sorted(row_data.keys())
        matrix = []
        for ry in sorted_group_ys:
            row_cells = []
            for cx in all_x0s:
                cell_block = row_data[ry].get(cx, None)
                row_cells.append(cell_block)
            matrix.append(row_cells)

        final_table_groups.append({
            "rows": matrix,
            "col_positions": all_x0s,
            "y_min": sorted_group_ys[0],
            "y_max": sorted_group_ys[-1],
        })

    return final_table_groups, non_table_blocks

# ─────────────────────────────────────────────
# Main extraction per page
# ─────────────────────────────────────────────

def extract_page_content(page, doc, translation_map=None):
    """
    Extract all layers from a single fitz page.
    Returns a dict with: geometry, text_blocks, images, table_groups
    """
    pw = page.rect.width
    ph = page.rect.height

    # ── Layer 2: Text spans with full formatting ──────────────
    # Use 'dict' mode — works reliably across all PyMuPDF versions
    raw = page.get_text("dict")

    text_blocks = []
    for block in raw.get("blocks", []):
        if block.get("type") != 0:  # 0 = text block
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                txt = span.get("text", "").strip()
                if not txt:
                    continue
                color = _color_tuple(span.get("color", 0))
                text_blocks.append({
                    "text": txt,
                    "x0": span["bbox"][0],
                    "y0": span["bbox"][1],
                    "x1": span["bbox"][2],
                    "y1": span["bbox"][3],
                    "font": span.get("font", "Arial"),
                    "size": span.get("size", 11.0),
                    "flags": span.get("flags", 0),
                    "color": color,
                    "bold": _is_bold(span.get("flags", 0), span.get("font", "")),
                    "italic": _is_italic(span.get("flags", 0), span.get("font", "")),
                })

    # Apply translation map: exact span match + substring match
    if translation_map:
        for tb in text_blocks:
            original = tb["text"]
            # Exact match first
            if original in translation_map:
                tb["text"] = translation_map[original]
            else:
                # Substring replacement for multi-word keys
                for orig_key, translated in translation_map.items():
                    if len(orig_key) > 3 and orig_key in tb["text"]:
                        tb["text"] = tb["text"].replace(orig_key, translated)

    # ── Layer 3: Embedded images ─────────────────────────────
    images = []
    seen_xrefs = set()
    for img_info in page.get_images(full=True):
        xref = img_info[0]
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)
        try:
            img_data = doc.extract_image(xref)
            if img_data and img_data.get("image"):
                # Get image placement bbox on the page
                img_list = page.get_image_bbox(img_info)
                if img_list:
                    rect = img_list
                else:
                    rect = fitz.Rect(0, 0, pw, ph)
                images.append({
                    "bytes": img_data["image"],
                    "ext": img_data.get("ext", "png"),
                    "x0": rect.x0 if hasattr(rect, 'x0') else 0,
                    "y0": rect.y0 if hasattr(rect, 'y0') else 0,
                    "x1": rect.x1 if hasattr(rect, 'x1') else pw,
                    "y1": rect.y1 if hasattr(rect, 'y1') else ph,
                    "width_pts": (rect.x1 - rect.x0) if hasattr(rect, 'x1') else pw,
                    "height_pts": (rect.y1 - rect.y0) if hasattr(rect, 'y1') else ph,
                })
        except Exception:
            pass

    # ── Layer 4: Table detection ─────────────────────────────
    table_groups, remaining_text_blocks = detect_tables(text_blocks, pw)

    return {
        "width_pts": pw,
        "height_pts": ph,
        "text_blocks": remaining_text_blocks,
        "images": images,
        "table_groups": table_groups,
    }

# ─────────────────────────────────────────────
# Word document builder
# ─────────────────────────────────────────────

MIN_GAP_COMP      = 0.30
MIN_FONT_SCALE    = 0.75
SAFETY_BUFFER     = 18.0

def estimate_height(blocks, font_scale, gap_comp, margin_top):
    physical_y = margin_top
    prev_y1 = margin_top
    for block in blocks:
        gap = max(0.0, block["y0"] - prev_y1) * gap_comp
        if block["type"] == "image":
            # Images are not scaled down by font scale
            row_h = block["y1"] - block["y0"]
        else:
            row_h = (block["y1"] - block["y0"]) * font_scale * 0.92
        physical_y += gap + row_h
        prev_y1 = block["y1"]
    return physical_y

def choose_compression(blocks, margin_top, margin_bottom, ph):
    budget = ph - margin_bottom - SAFETY_BUFFER

    # Quick exit: content already fits at 1.0 / 1.0
    if estimate_height(blocks, 1.0, 1.0, margin_top) <= budget:
        return 1.0, 1.0, False

    # Step A: compress gaps only
    for gc in [0.80, 0.60, 0.45, 0.30]:
        if estimate_height(blocks, 1.0, gc, margin_top) <= budget:
            return 1.0, gc, False

    # Step B: compress gaps at floor + shrink fonts
    for fs in [0.95, 0.90, 0.85, 0.80, 0.75]:
        if estimate_height(blocks, fs, MIN_GAP_COMP, margin_top) <= budget:
            return fs, MIN_GAP_COMP, False

    # Step C: progressive extra compression to avoid page ballooning if possible
    for fs in [0.70, 0.65]:
        if estimate_height(blocks, fs, 0.20, margin_top) <= budget:
            return fs, 0.20, False

    # Cannot fit — signal for page splitting
    return MIN_FONT_SCALE, MIN_GAP_COMP, True

def split_blocks_to_fit(blocks, margin_top, margin_bottom, ph):
    budget = ph - margin_bottom - SAFETY_BUFFER
    chunks = []
    current_chunk = []
    font_scale, gap_comp = MIN_FONT_SCALE, MIN_GAP_COMP

    for block in blocks:
        test = current_chunk + [block]
        h = estimate_height(test, font_scale, gap_comp, margin_top)
        if h <= budget or not current_chunk:
            current_chunk.append(block)
        else:
            chunks.append(current_chunk)
            current_chunk = [block]

    if current_chunk:
        chunks.append(current_chunk)

    return chunks

def _apply_run_formatting(run, tb, font_scale=1.0):
    """Apply font name, size, bold, italic, color to a run."""
    # Clean up font name to a safe Word font
    font_name = tb.get("font", "Arial")
    # Strip encoding suffixes like +Arial, ABCDEF+Arial → Arial
    font_name = re.sub(r'^[A-Z]{6}\+', '', font_name)
    font_name = re.sub(r'[,\-](Bold|Italic|Regular|Light|Medium|Black|Heavy).*$', '', font_name, flags=re.IGNORECASE)
    # Map common PDF font names to Word-safe equivalents
    font_map = {
        "TimesNewRoman": "Times New Roman",
        "CourierNew": "Courier New",
        "HelveticaNeue": "Arial",
        "Helvetica": "Arial",
        "Calibri": "Calibri",
        "Garamond": "Garamond",
    }
    for pdf_name, word_name in font_map.items():
        if pdf_name.lower() in font_name.lower():
            font_name = word_name
            break

    try:
        run.font.name = font_name
    except Exception:
        run.font.name = "Arial"
    
    # Scale font size
    base_size = tb.get("size", 11.0)
    run.font.size = Pt(max(1.0, base_size * font_scale))
    run.bold = tb.get("bold", False)
    run.italic = tb.get("italic", False)
    run.font.color.rgb = tb.get("color", RGBColor(0, 0, 0))

def build_docx(pages_content, output_path):
    """Assemble the final Word document from extracted page content."""
    doc = Document()

    # Remove default empty paragraph in fresh document
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    for page_idx, page_content in enumerate(pages_content):
        pw = page_content["width_pts"]
        ph = page_content["height_pts"]
        text_blocks = page_content["text_blocks"]
        images = page_content["images"]
        table_groups = page_content["table_groups"]

        # Compute content margins from the earliest/latest positions of all elements
        all_x0s = []
        all_y0s = []
        all_x1s = []
        all_y1s = []
        for tb in text_blocks:
            all_x0s.append(tb["x0"])
            all_y0s.append(tb["y0"])
            all_x1s.append(tb["x1"])
            all_y1s.append(tb["y1"])
        for tg in table_groups:
            for row in tg["rows"]:
                for cell_block in row:
                    if cell_block:
                        all_x0s.append(cell_block["x0"])
                        all_y0s.append(cell_block["y0"])
                        all_x1s.append(cell_block["x1"])
                        all_y1s.append(cell_block["y1"])
        for img in images:
            all_x0s.append(img["x0"])
            all_y0s.append(img["y0"])
            all_x1s.append(img["x1"])
            all_y1s.append(img["y1"])

        if all_x0s:
            min_x = min(all_x0s)
            min_y = min(all_y0s)
            max_x = max(all_x1s)
            max_y = max(all_y1s)
            left_margin   = max(18.0, min_x - 5.0)
            top_margin    = max(18.0, min_y - 5.0)
            right_margin  = max(18.0, pw - max_x - 5.0)
            bottom_margin = max(18.0, ph - max_y - 5.0)
        else:
            left_margin = top_margin = right_margin = bottom_margin = 36.0

        # Group text blocks into paragraph lines
        # Sort by Y then X
        sorted_text_blocks = sorted(text_blocks, key=lambda b: (round(b["y0"] / 4) * 4, b["x0"]))
        line_groups = []
        current_line = []
        last_y = None
        for tb in sorted_text_blocks:
            if last_y is None or abs(tb["y0"] - last_y) < 5.0:
                current_line.append(tb)
            else:
                if current_line:
                    line_groups.append(current_line)
                current_line = [tb]
            last_y = tb["y0"]
        if current_line:
            line_groups.append(current_line)

        # Build unified blocks list
        blocks = []
        for img in images:
            blocks.append({
                "type": "image",
                "y0": img["y0"],
                "y1": img["y1"],
                "data": img
            })
        for tg in table_groups:
            blocks.append({
                "type": "table",
                "y0": tg["y_min"],
                "y1": tg["y_max"],
                "data": tg
            })
        for line in line_groups:
            line_y0 = min(tb["y0"] for tb in line)
            line_y1 = max(tb["y1"] for tb in line)
            blocks.append({
                "type": "paragraph",
                "y0": line_y0,
                "y1": line_y1,
                "data": line
            })

        # Sort blocks by top Y coordinate
        blocks.sort(key=lambda b: b["y0"])

        # Choose compression
        font_scale, gap_comp, overflows = choose_compression(
            blocks, top_margin, bottom_margin, ph
        )

        if overflows:
            sub_pages = split_blocks_to_fit(blocks, top_margin, bottom_margin, ph)
        else:
            sub_pages = [blocks]

        # Render sub-pages as sections
        for sub_idx, sub_blocks in enumerate(sub_pages):
            if page_idx == 0 and sub_idx == 0:
                section = doc.sections[0]
            else:
                section = doc.add_section()

            section.page_width = Pt(pw)
            section.page_height = Pt(ph)
            section.orientation = WD_ORIENT.LANDSCAPE if pw > ph else WD_ORIENT.PORTRAIT
            section.left_margin = Pt(left_margin)
            section.right_margin = Pt(right_margin)
            section.top_margin = Pt(top_margin)
            section.bottom_margin = Pt(bottom_margin)

            prev_y1 = top_margin

            for block in sub_blocks:
                btype = block["type"]
                data = block["data"]
                gap = max(0.0, block["y0"] - prev_y1) * gap_comp

                if btype == "image":
                    img = data
                    # Check gap
                    if gap > 0.5:
                        sp = doc.add_paragraph()
                        sp.paragraph_format.space_before = Pt(gap)
                        sp.paragraph_format.space_after = Pt(0)
                        sp.paragraph_format.line_spacing = Pt(1)
                        sp.add_run().font.size = Pt(1)

                    img_para = doc.add_paragraph()
                    img_para.paragraph_format.space_before = Pt(0)
                    img_para.paragraph_format.space_after = Pt(0)
                    img_run = img_para.add_run()
                    img_buf = io.BytesIO(img["bytes"])
                    try:
                        if img.get("is_full_page"):
                            content_w_in = (pw - left_margin - right_margin) / 72
                            content_h_in = (ph - top_margin - bottom_margin) / 72
                            img_run.add_picture(img_buf, width=Inches(max(0.5, content_w_in)), height=Inches(max(0.5, content_h_in)))
                        else:
                            width_in = img["width_pts"] / 72
                            height_in = img["height_pts"] / 72
                            content_w = (pw - left_margin - right_margin) / 72
                            if width_in > content_w:
                                scale = content_w / width_in
                                width_in *= scale
                                height_in *= scale
                            img_run.add_picture(img_buf, width=Inches(width_in), height=Inches(height_in))
                    except Exception:
                        try:
                            img_run.add_picture(img_buf)
                        except Exception:
                            pass
                    prev_y1 = block["y1"]

                elif btype == "table":
                    table_group = data
                    matrix = table_group["rows"]
                    col_positions = table_group["col_positions"]
                    n_cols = len(col_positions)
                    n_rows = len(matrix)
                    if n_cols < 1 or n_rows < 1:
                        continue
                    if gap > 5:
                        # Estimate spaces representing the horizontal gap
                        font_size = prev.get("size", 10.0)
                        char_width = max(2.5, font_size * 0.25)
                        num_spaces = max(1, int(gap / char_width))
                        space_run = para.add_run(" " * num_spaces)
                        _apply_run_formatting(space_run, prev)
                run = para.add_run(text)
                _apply_run_formatting(run, tb)

    doc.save(output_path)
    print(f"✅ Successfully saved layered Word document to: {output_path}")

# ─────────────────────────────────────────────
# Public entry point (callable from backends)
# ─────────────────────────────────────────────

def convert_pdf_to_docx(pdf_path, output_docx_path, translation_map=None):
    """
    Main function: extract all layers from PDF, build editable Word document.

    Handles three types of pages automatically:
      1. Native text PDF  → Extract text spans, tables, images → editable Word
      2. Scanned/image PDF → Render page as image + insert into Word (readable but not editable)
      3. Mixed PDF        → Text pages get editable text; image pages get rendered images

    translation_map: dict {original_text: translated_text}
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    fitz_doc = fitz.open(pdf_path)
    
    # Check if PDF is scanned or image-only
    scanned_pages = 0
    for page_idx in range(len(fitz_doc)):
        page = fitz_doc[page_idx]
        pw, ph = page.rect.width, page.rect.height
        text = page.get_text().strip()
        images = page.get_images(full=True)
        
        # 1. No text
        if not text:
            scanned_pages += 1
            continue
        # 2. Sparse text with images
        if len(text) < 150 and images:
            scanned_pages += 1
            continue
        # 3. Giant image covering > 85% of the page
        is_page_scanned = False
        for img_info in images:
            try:
                img_rects = page.get_image_bbox(img_info)
                rect = img_rects if isinstance(img_rects, fitz.Rect) else (img_rects[0] if img_rects else None)
                if rect:
                    img_w = rect.width
                    img_h = rect.height
                    if (img_w * img_h) > (pw * ph * 0.85):
                        if len(text) < 1200 or len(images) == 1:
                            is_page_scanned = True
                            break
            except:
                pass
        if is_page_scanned:
            scanned_pages += 1
            continue

    if scanned_pages > 0 or len(fitz_doc) == 0:
        print("Error: Scanned pages detected. This PDF requires OCR.")
        sys.exit(2)

    pages_content = []

    for page_idx in range(len(fitz_doc)):
        page = fitz_doc[page_idx]
        pw = page.rect.width
        ph = page.rect.height

        # ── Detect if this page is scanned (image-only) ──
        raw_text = page.get_text("text").strip()
        images_on_page = page.get_images(full=True)

        # A page is considered scanned if:
        # - Very little extractable text (< 30 chars), AND
        # - Has at least one embedded image that covers most of the page
        is_scanned = False
        if len(raw_text) < 30 and images_on_page:
            for img_info in images_on_page:
                try:
                    img_rect = page.get_image_bbox(img_info)
                    img_area = (img_rect.x1 - img_rect.x0) * (img_rect.y1 - img_rect.y0)
                    page_area = pw * ph
                    if img_area > page_area * 0.5:
                        is_scanned = True
                        break
                except Exception:
                    pass
            # Also treat as scanned if literally no text at all
            if len(raw_text) == 0:
                is_scanned = True

        if is_scanned:
            # ── Scanned page: render to high-res image and embed in Word ──
            # Render at 2x resolution for clarity
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            pages_content.append({
                "width_pts": pw,
                "height_pts": ph,
                "text_blocks": [],
                "images": [{
                    "bytes": img_bytes,
                    "ext": "png",
                    "x0": 0, "y0": 0,
                    "x1": pw, "y1": ph,
                    "width_pts": pw,
                    "height_pts": ph,
                    "is_full_page": True,
                }],
                "table_groups": [],
                "is_scanned": True,
            })
            print(f"  Page {page_idx + 1}: scanned image — rendered as full-page image")
        else:
            # ── Native text page: full layer extraction ──
            content = extract_page_content(page, fitz_doc, translation_map)
            content["is_scanned"] = False
            pages_content.append(content)
            print(f"  Page {page_idx + 1}: native text — {len(content['text_blocks'])} spans, {len(content['table_groups'])} tables, {len(content['images'])} images")

    fitz_doc.close()
    build_docx(pages_content, output_docx_path)
    return output_docx_path



# ─────────────────────────────────────────────
# CLI entry point
# ─────────────────────────────────────────────

def main():
    if len(sys.argv) < 3:
        print("Usage: pdf_layer_extractor.py <input_pdf> <output_docx> [<translation_json>]")
        sys.exit(1)

    pdf_path = sys.argv[1]
    output_docx_path = sys.argv[2]
    translation_map = None

    if len(sys.argv) >= 4:
        tx_path = sys.argv[3]
        if os.path.exists(tx_path):
            with open(tx_path, "r", encoding="utf-8") as f:
                translation_map = json.load(f)

    try:
        convert_pdf_to_docx(pdf_path, output_docx_path, translation_map)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
