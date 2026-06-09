import os
import io
import re
import json
import httpx
import asyncio
import tempfile
import smtplib
import base64
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, Request, Response, HTTPException, Query, UploadFile, File
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pdf2docx import Converter
import fitz

app = FastAPI(title="EPS Translator Backend", description="FastAPI Backend for Document translation and layout analysis")

# Enable CORS for frontend integration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Configuration (set via environment variables) ---
# GEMINI_API_KEY, CLOUDCONVERT_API_KEY, GOOGLE_SHEETS_URL, PROPOSAL_DESTINATION_EMAIL

# --- Helper Functions for Text Normalization & Layout Resolution ---
def get_language_code(lang: str) -> str:
    lang = lang.lower().strip()
    mapping = {
        'french': 'fr', 'spanish': 'es', 'german': 'de', 'italian': 'it',
        'portuguese': 'pt', 'russian': 'ru', 'chinese': 'zh', 'japanese': 'ja',
        'korean': 'ko', 'arabic': 'ar', 'hindi': 'hi', 'dutch': 'nl',
        'swedish': 'sv', 'polish': 'pl', 'turkish': 'tr', 'vietnamese': 'vi',
        'english': 'en', 'danish': 'da', 'finnish': 'fi', 'norwegian': 'no',
        'czech': 'cs', 'greek': 'el', 'romanian': 'ro', 'hungarian': 'hu',
        'indonesian': 'id', 'thai': 'th', 'ukrainian': 'uk'
    }
    return mapping.get(lang, lang[:2])

def normalize_text(text: str) -> str:
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text).strip()

def iter_block_paragraphs(parent):
    if hasattr(parent, 'paragraphs'):
        for p in parent.paragraphs:
            yield p
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_paragraphs(cell)

def adjust_table_row_heights(parent):
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            for row in table.rows:
                trPr = row._tr.get_or_add_trPr()
                trHeight = trPr.first_child_found_in("w:trHeight")
                if trHeight is not None:
                    hRule = trHeight.get(qn('w:hRule'))
                    if hRule == 'exact':
                        trHeight.set(qn('w:hRule'), 'atLeast')
                for cell in row.cells:
                    adjust_table_row_heights(cell)

def replace_text_in_paragraph_single(p, key: str, replacement: str) -> bool:
    if key not in p.text:
        return False

    len_orig = len(key)
    len_trans = len(replacement)
    scale_factor = 1.0
    if len_orig > 5 and len_trans > len_orig:
        scale_factor = max(0.70, len_orig / len_trans)

    def scale_run(run_obj):
        if scale_factor < 1.0 and run_obj.font.size:
            run_obj.font.size = Pt(run_obj.font.size.pt * scale_factor)

    if p.text.strip() == key.strip():
        if p.runs:
            for r in p.runs:
                scale_run(r)
            p.runs[0].text = replacement
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = replacement
        return True

    if not p.runs:
        p.text = p.text.replace(key, replacement)
        return True

    for r in p.runs:
        if key in r.text:
            scale_run(r)
            r.text = r.text.replace(key, replacement)
            return True

    char_map = []
    for r_idx, r in enumerate(p.runs):
        if r.text:
            for c_idx in range(len(r.text)):
                char_map.append((r_idx, c_idx))

    if len(char_map) != len(p.text):
        p.text = p.text.replace(key, replacement)
        return True

    start_idx = p.text.find(key)
    if start_idx == -1:
        return False

    end_idx = start_idx + len(key)
    involved_runs = sorted(list(set(char_map[i][0] for i in range(start_idx, end_idx))))
    if not involved_runs:
        p.text = p.text.replace(key, replacement)
        return True

    first_r_idx, first_c_idx = char_map[start_idx]
    last_r_idx, last_c_idx = char_map[end_idx - 1]

    first_run = p.runs[first_r_idx]
    last_run = p.runs[last_r_idx]

    prefix = first_run.text[:first_c_idx]
    suffix = last_run.text[last_c_idx + 1:]

    for r_idx in involved_runs:
        scale_run(p.runs[r_idx])

    first_run.text = prefix + replacement
    for r_idx in involved_runs[1:]:
        p.runs[r_idx].text = ""

    if first_r_idx != last_r_idx:
        last_run.text = suffix
    else:
        first_run.text += suffix

    return True

def replace_text_in_paragraph(p, key: str, replacement: str) -> bool:
    limit = 10
    replaced = False
    while key in p.text and limit > 0:
        if replace_text_in_paragraph_single(p, key, replacement):
            replaced = True
        limit -= 1
    return replaced

def apply_docx_translation(docx_path: str, output_path: str, translation_map: dict, autofit: bool = False):
    doc = Document(docx_path)

    # Change exact height row rules to atLeast to avoid vertical text clipping
    adjust_table_row_heights(doc)
    for section in doc.sections:
        if section.header:
            adjust_table_row_heights(section.header)
        if section.footer:
            adjust_table_row_heights(section.footer)

    if autofit:
        for table in doc.tables:
            table.autofit = True
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcW = tcPr.first_child_found_in("w:tcW")
                    if tcW is not None:
                        tcW.set(qn('w:w'), '0')
                        tcW.set(qn('w:type'), 'auto')

    normalized_map = {}
    original_map = {}
    
    for original, translated in translation_map.items():
        if not original or not translated:
            continue
        orig_cleaned = original.strip()
        trans_cleaned = translated.strip()
        if not orig_cleaned:
            continue
            
        original_map[orig_cleaned] = trans_cleaned
        norm_orig = normalize_text(orig_cleaned)
        if norm_orig:
            normalized_map[norm_orig] = trans_cleaned

    sorted_norm_keys = sorted(normalized_map.keys(), key=len, reverse=True)
    sorted_orig_keys = sorted(original_map.keys(), key=len, reverse=True)

    paragraphs = []
    for p in iter_block_paragraphs(doc):
        paragraphs.append(p)
        
    for section in doc.sections:
        if section.header:
            for p in iter_block_paragraphs(section.header):
                paragraphs.append(p)
        if section.footer:
            for p in iter_block_paragraphs(section.footer):
                paragraphs.append(p)

    for p in paragraphs:
        if not p.text.strip():
            continue
            
        p_text_norm = normalize_text(p.text)
        replaced_p = False

        # Phase A: Exact paragraph match on normalized text
        for key in sorted_norm_keys:
            if p_text_norm == key:
                replacement = normalized_map[key]
                len_orig = len(key)
                len_trans = len(replacement)
                scale_factor = 1.0
                if len_orig > 5 and len_trans > len_orig:
                    scale_factor = max(0.70, len_orig / len_trans)

                if p.runs:
                    for r in p.runs:
                        if scale_factor < 1.0 and r.font.size:
                            r.font.size = Pt(r.font.size.pt * scale_factor)
                    p.runs[0].text = replacement
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = replacement
                replaced_p = True
                break

        if replaced_p:
            continue

        # Phase B: Substring match on original keys (un-normalized)
        for orig_key in sorted_orig_keys:
            if orig_key in p.text:
                replacement = original_map[orig_key]
                replace_text_in_paragraph(p, orig_key, replacement)

    doc.save(output_path)

# --- Translation Logic (MyMemory + Gemini fallbacks) ---
async def translate_mymemory(texts: List[str], target_lang: str, source_lang: str = 'English') -> Optional[Dict[str, str]]:
    lang_pair = f"{get_language_code(source_lang)}|{get_language_code(target_lang)}"
    results = {}
    
    async with httpx.AsyncClient(timeout=10.0) as client:
        tasks = []
        for text in texts:
            trimmed = text.strip()
            if not trimmed or trimmed.isnumeric():
                results[text] = text
                continue
            if re.match(r'^\d+(\.\d+)?\s*(mm|cm|m|in|inch|inches|px|pt|%|°|deg|rad)?$', trimmed, re.I):
                results[text] = text
                continue
                
            tasks.append((text, client.get("https://api.mymemory.translated.net/get", params={"q": text, "langpair": lang_pair})))
            
        if not tasks:
            return results
            
        responses = await asyncio.gather(*(t[1] for t in tasks), return_exceptions=True)
        for (text, _), response in zip(tasks, responses):
            if isinstance(response, Exception) or response.status_code != 200:
                return None
            data = response.json()
            if data.get("responseStatus") == 200:
                results[text] = data.get("responseData", {}).get("translatedText", text)
            else:
                return None
    return results

async def translate_gemini(texts: List[str], target_lang: str, source_lang: str = 'English') -> Optional[Dict[str, str]]:
    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    if not gemini_key:
        return None

    prompt_text = (
        f"You are an expert technical translator specialising in engineering drawings.\n"
        f"Translate text labels from {source_lang} into {target_lang}.\n"
        f"Keep technical numbers, CURRENT current ratings (e.g. 10A, 15A), wire codes, and standard abbreviations (GND, VCC, ECU) EXACTLY as-is.\n"
        f"Preserve capitalization styles (e.g. ALL CAPS).\n"
        f"Return a clean JSON object mapping original text to translated text. No markdown fences.\n"
        f"Input texts: {json.dumps(texts, ensure_ascii=False)}"
    )

    payload = {
        "contents": [{"parts": [{"text": prompt_text}]}],
        "generationConfig": {
            "responseMimeType": "application/json"
        }
    }

    models = ["gemini-2.5-flash", "gemini-2.5-flash-lite", "gemini-2.5-pro"]
    async with httpx.AsyncClient(timeout=30.0) as client:
        for model in models:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={gemini_key}"
            try:
                r = await client.post(url, json=payload)
                if r.status_code == 200:
                    data = r.json()
                    resp_text = data['candidates'][0]['content']['parts'][0]['text']
                    cleaned = resp_text.strip().removeprefix("```json").removesuffix("```").strip()
                    return json.loads(cleaned)
            except Exception as e:
                print(f"Gemini {model} failed: {e}")
    return None

# --- CloudConvert Async Connector ---
def get_cloudconvert_ocr_lang(lang: str) -> Optional[str]:
    if not lang:
        return None
    lang = lang.lower().strip()
    mapping = {
        'english': 'eng', 'german': 'deu', 'french': 'fra', 'spanish': 'spa',
        'italian': 'ita', 'portuguese': 'por', 'russian': 'rus', 'chinese': 'chi',
        'japanese': 'jpn', 'korean': 'kor', 'dutch': 'nld', 'swedish': 'swe',
        'polish': 'pol', 'turkish': 'tur'
    }
    return mapping.get(lang, None)

def is_scanned_pdf(pdf_path: str) -> bool:
    try:
        doc = fitz.open(pdf_path)
        total_text = ""
        for page in doc:
            total_text += page.get_text()
            if len(total_text.strip()) > 30:
                return False
        return True
    except Exception as e:
        print(f"Error checking if PDF is scanned: {e}")
        return True # Default to true as safer fallback

async def run_cloudconvert_job(input_path: str, from_format: str, to_format: str, api_key: str, ocr: bool = False, ocr_language: str = None) -> str:
    async with httpx.AsyncClient(timeout=30.0) as client:
        headers = {"Authorization": f"Bearer {api_key}"}
        
        convert_task = {
            "operation": "convert",
            "input": ["import-file"],
            "input_format": from_format,
            "output_format": to_format
        }
        if ocr:
            convert_task["ocr"] = True
            if ocr_language:
                convert_task["ocr_language"] = ocr_language

        # 1. Create Job
        payload = {
            "tasks": {
                "import-file": {
                    "operation": "import/upload"
                },
                "convert-file": convert_task,
                "export-file": {
                    "operation": "export/url",
                    "input": ["convert-file"]
                }
            }
        }
        
        r = await client.post("https://api.cloudconvert.com/v2/jobs", json=payload, headers=headers)
        if r.status_code != 201:
            raise HTTPException(status_code=500, detail=f"CloudConvert create job failed: {r.text}")
        
        job = r.json()
        job_id = job['data']['id']
        
        upload_task = next(t for t in job['data']['tasks'] if t['name'] == 'import-file')
        upload_url = upload_task['result']['form']['url']
        upload_params = upload_task['result']['form']['parameters']
        
        # 2. Upload file
        with open(input_path, 'rb') as f:
            files = {'file': f}
            up_r = await client.post(upload_url, data=upload_params, files=files)
            if up_r.status_code not in [201, 204, 200]:
                raise HTTPException(status_code=500, detail=f"CloudConvert upload failed: {up_r.text}")
                
        # 3. Poll for status
        for _ in range(30): # 90 seconds timeout
            await asyncio.sleep(3.0)
            status_r = await client.get(f"https://api.cloudconvert.com/v2/jobs/{job_id}", headers=headers)
            if status_r.status_code != 200:
                continue
            status_job = status_r.json()
            job_status = status_job['data']['status']
            
            if job_status == 'finished':
                export_task = next(t for t in status_job['data']['tasks'] if t['name'] == 'export-file')
                export_url = export_task['result']['files'][0]['url']
                return export_url
            elif job_status == 'failed':
                raise HTTPException(status_code=500, detail="CloudConvert conversion job failed on server.")
                
        raise HTTPException(status_code=500, detail="CloudConvert job execution timed out.")

# --- API Data Models ---
class OcrExtractRequest(BaseModel):
    image: str
    sourceLanguage: str = "German"

class TextboxModel(BaseModel):
    text: str
    x: float
    y: float
    w: float
    h: float
    fsFraction: float
    bold: bool

class PageModel(BaseModel):
    width: float
    height: float
    background: str
    textboxes: List[TextboxModel]

class PdfToDocxRequest(BaseModel):
    originalPdfBase64: Optional[str] = ""
    translationMap: Optional[dict] = None
    filename: Optional[str] = "document"
    cloudConvertApiKey: Optional[str] = ""
    sourceLanguage: Optional[str] = "German"
    isLayoutReplica: Optional[bool] = False
    pages: Optional[List[PageModel]] = None

class AiTranslateRequest(BaseModel):
    texts: List[str]
    targetLanguage: str
    sourceLanguage: str = "English"

class LogLoginRequest(BaseModel):
    name: str
    email: str

class SendProposalRequest(BaseModel):
    name: str
    email: str
    filename: str
    targetLanguage: str
    sourceLanguage: str = "English"
    svgText: str
    message: Optional[str] = ""

# --- REST Routes ---

@app.get("/api/check-env")
async def check_env():
    gemini_configured = len(os.environ.get("GEMINI_API_KEY", "")) > 15
    cc_configured = len(os.environ.get("CLOUDCONVERT_API_KEY", "")) > 15
    return {
        "status": "ready",
        "gemini_configured": gemini_configured,
        "cloudconvert_configured": cc_configured,
        "environment": "production" if "DYNO" in os.environ or "RAILWAY_STATIC_URL" in os.environ else "development"
    }

@app.post("/api/log-login")
async def log_login(req: LogLoginRequest):
    sheets_url = os.environ.get("GOOGLE_SHEETS_URL", "")
    payload = {"name": req.name, "email": req.email, "timestamp": "now"}
    async with httpx.AsyncClient() as client:
        try:
            r = await client.post(sheets_url, json=payload)
            return {"success": r.status_code == 200}
        except Exception as e:
            return {"success": False, "error": str(e)}

@app.post("/api/ai-translate")
async def ai_translate(req: AiTranslateRequest):
    # Try MyMemory API first
    tx_map = await translate_mymemory(req.texts, req.targetLanguage, req.sourceLanguage)
    if tx_map is None:
        # Fall back to Gemini API
        tx_map = await translate_gemini(req.texts, req.targetLanguage, req.sourceLanguage)
        
    if tx_map is not None:
        return tx_map
        
    raise HTTPException(status_code=500, detail="Translation service failed. Please try again.")

@app.post("/api/ocr-extract")
async def ocr_extract(req: OcrExtractRequest):
    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    if not gemini_key:
        raise HTTPException(status_code=500, detail="Gemini API key is not configured on the backend.")

    raw_image = req.image
    if raw_image.startswith("data:image"):
        raw_image = raw_image.split(",")[1]

    prompt_text = (
        "You are an expert OCR tool specializing in layout documents.\n"
        f"Identify all visible text labels written in {req.sourceLanguage}.\n"
        "Locate their bounding box coordinates on the image using a scale of 0 to 1000 (where 0,0 is top-left and 1000,1000 is bottom-right).\n\n"
        "Return a JSON array of objects. Each object must contain exactly:\n"
        "- 'text': The detected text segment\n"
        "- 'box': An object with 'ymin', 'xmin', 'ymax', 'xmax' as integers between 0 and 1000.\n\n"
        "Example format:\n"
        "[\n"
        "  {\n"
        "    \"text\": \"HAUPTSCHALTER\",\n"
        "    \"box\": { \"ymin\": 120, \"xmin\": 300, \"ymax\": 145, \"xmax\": 450 }\n"
        "  }\n"
        "]"
    )

    payload = {
        "contents": [{
            "parts": [
                {"text": prompt_text},
                {"inlineData": {"mimeType": "image/jpeg", "data": raw_image}}
            ]
        }],
        "generationConfig": {
            "responseMimeType": "application/json"
        }
    }

    async with httpx.AsyncClient(timeout=45.0) as client:
        models = ["gemini-2.5-flash", "gemini-2.5-flash-lite", "gemini-2.5-pro"]
        for model in models:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={gemini_key}"
            try:
                r = await client.post(url, json=payload)
                if r.status_code == 200:
                    data = r.json()
                    resp_text = data['candidates'][0]['content']['parts'][0]['text']
                    cleaned = resp_text.strip().removeprefix("```json").removesuffix("```").strip()
                    return JSONResponse(content=json.loads(cleaned))
                elif r.status_code == 429:
                    continue # Try next model or fail
            except Exception as e:
                print(f"OCR Gemini {model} failed: {e}")
                
        raise HTTPException(status_code=429, detail="The OCR service limit was reached. Please try again.")

@app.post("/api/pdf-to-docx")
async def pdf_to_docx(req: PdfToDocxRequest):
    import time
    import html
    from docx import Document
    from docx.shared import Pt
    from docx.enum.section import WD_ORIENT
    from docx.oxml import parse_xml

    filename_stem = req.filename or "translated"
    # Sanitize filename for HTTP headers (ASCII only)
    import unicodedata
    safe_filename = unicodedata.normalize('NFKD', filename_stem).encode('ascii', 'ignore').decode('ascii')
    safe_filename = re.sub(r'[^a-zA-Z0-9_\-.]', '_', safe_filename) or 'document'

    # --- 1. EXACT REPLICA CONVERSION (via backgrounds & absolute textboxes) ---
    if req.isLayoutReplica and req.pages:
        try:
            doc = Document()
            
            for index, page in enumerate(req.pages):
                # Add a section for each page (except the first one which is created by default)
                if index == 0:
                    section = doc.sections[0]
                else:
                    section = doc.add_section()
                    section.header.is_linked_to_previous = False
                
                # Set orientation and page size
                # Width and height are in points (72 points per inch)
                width_pt = page.width
                height_pt = page.height
                
                section.page_width = Pt(width_pt)
                section.page_height = Pt(height_pt)
                
                if width_pt > height_pt:
                    section.orientation = WD_ORIENT.LANDSCAPE
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                
                # Set margins to 0 for full bleed background and absolute positioning
                section.top_margin = Pt(0)
                section.bottom_margin = Pt(0)
                section.left_margin = Pt(0)
                section.right_margin = Pt(0)
                
                # Setup background image in the section header
                if page.background:
                    bg_bytes = base64.b64decode(page.background)
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_bg:
                        tmp_bg.write(bg_bytes)
                        tmp_bg_path = tmp_bg.name
                    
                    try:
                        # Add image relationship to the header part
                        header = section.header
                        rId, _ = header.part.get_or_add_image(tmp_bg_path)
                        
                        # Calculate size in EMUs
                        w_emus = int(width_pt * 12700)
                        h_emus = int(height_pt * 12700)
                        bg_id = int(time.time() * 1000 + index) % 1000000
                        
                        bg_drawing_xml = f"""
                        <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                                   xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                                   xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                                   xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
                                   xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
                            <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
                                <wp:simplePos x="0" y="0"/>
                                <wp:positionH relativeFrom="page">
                                    <wp:posOffset>0</wp:posOffset>
                                </wp:positionH>
                                <wp:positionV relativeFrom="page">
                                    <wp:posOffset>0</wp:posOffset>
                                </wp:positionV>
                                <wp:extent cx="{w_emus}" cy="{h_emus}"/>
                                <wp:docPr id="{bg_id}" name="Bg_{bg_id}"/>
                                <wp:cNvGraphicFramePr/>
                                <a:graphic>
                                    <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                                        <pic:pic>
                                            <pic:nvPicPr>
                                                <pic:cNvPr id="{bg_id}" name="BgPic_{bg_id}"/>
                                                <pic:cNvPicPr/>
                                            </pic:nvPicPr>
                                            <pic:blipFill>
                                                <a:blip r:embed="{rId}"/>
                                                <a:stretch>
                                                    <a:fillRect/>
                                                </a:stretch>
                                            </pic:blipFill>
                                            <pic:spPr>
                                                <a:xfrm>
                                                    <a:off x="0" y="0"/>
                                                    <a:ext cx="{w_emus}" cy="{h_emus}"/>
                                                </a:xfrm>
                                                <a:prstGeom prst="rect">
                                                    <a:avLst/>
                                                </a:prstGeom>
                                            </pic:spPr>
                                        </pic:pic>
                                    </a:graphicData>
                                </a:graphic>
                            </wp:anchor>
                        </w:drawing>
                        """
                        # Insert in header paragraph run
                        h_p = header.paragraphs[0]
                        h_run = h_p.add_run()
                        h_run._r.append(parse_xml(bg_drawing_xml))
                    finally:
                        try:
                            os.unlink(tmp_bg_path)
                        except:
                            pass
                
                # Now, add body textboxes
                # Access or create the body paragraph
                body_p = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
                
                # Add textboxes
                for t_idx, box in enumerate(page.textboxes):
                    # Convert fraction coordinates to point dimensions
                    x_pt = box.x * width_pt
                    y_pt = box.y * height_pt
                    w_pt = box.w * width_pt
                    h_pt = box.h * height_pt
                    
                    x_emus = int(x_pt * 12700)
                    y_emus = int(y_pt * 12700)
                    w_emus = int(w_pt * 12700)
                    h_emus = int(h_pt * 12700)
                    
                    fs_pt = max(6.0, min(72.0, box.fsFraction * height_pt))
                    is_bold = box.bold
                    text_escaped = html.escape(box.text)
                    
                    tb_id = int(time.time() * 1000 + index * 1000 + t_idx) % 1000000
                    
                    textbox_xml = f"""
                    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                               xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
                            <wp:simplePos x="0" y="0"/>
                            <wp:positionH relativeFrom="page">
                                <wp:posOffset>{x_emus}</wp:posOffset>
                            </wp:positionH>
                            <wp:positionV relativeFrom="page">
                                <wp:posOffset>{y_emus}</wp:posOffset>
                            </wp:positionV>
                            <wp:extent cx="{w_emus}" cy="{h_emus}"/>
                            <wp:docPr id="{tb_id}" name="Text Box {tb_id}"/>
                            <wp:cNvGraphicFramePr/>
                            <a:graphic>
                                <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                                    <wps:wsp>
                                        <wps:spPr>
                                            <a:xfrm>
                                                <a:off x="0" y="0"/>
                                                <a:ext cx="{w_emus}" cy="{h_emus}"/>
                                            </a:xfrm>
                                            <a:prstGeom prst="rect">
                                                <a:avLst/>
                                            </a:prstGeom>
                                            <a:solidFill>
                                                <a:noFill/>
                                            </a:solidFill>
                                            <a:ln>
                                                <a:noFill/>
                                            </a:ln>
                                        </wps:spPr>
                                        <wps:txbx>
                                            <w:txbxContent>
                                                <w:p>
                                                    <w:pPr>
                                                        <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
                                                        <w:ind w:left="0" w:right="0"/>
                                                    </w:pPr>
                                                    <w:r>
                                                        <w:rPr>
                                                            <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                                                            <w:sz w:val="{int(fs_pt * 2)}"/>
                                                            {"<w:b/>" if is_bold else ""}
                                                        </w:rPr>
                                                        <w:t>{text_escaped}</w:t>
                                                    </w:r>
                                                </w:p>
                                            </w:txbxContent>
                                        </wps:txbx>
                                    </wps:wsp>
                                </a:graphicData>
                            </a:graphic>
                        </wp:anchor>
                    </w:drawing>
                    """
                    run = body_p.add_run()
                    run._r.append(parse_xml(textbox_xml))

            # Save generated replica to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            output_bytes = docx_buffer.read()

            return StreamingResponse(
                io.BytesIO(output_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={safe_filename}_translated.docx"}
            )
        except Exception as e:
            print(f"Replica document generation failed: {e}")
            # fall back automatically to standard flow if layout replica generation fails
            pass

    # --- 2. PDF → DOCX CONVERSION ---
    pdf_bytes = base64.b64decode(req.originalPdfBase64 or "")
    if not pdf_bytes:
        raise HTTPException(status_code=400, detail="Invalid or missing PDF data.")

    tx_map = req.translationMap or {}

    with tempfile.TemporaryDirectory() as tmpdir:
        input_pdf = os.path.join(tmpdir, "input.pdf")
        clean_docx = os.path.join(tmpdir, "clean.docx")
        final_docx = os.path.join(tmpdir, "final.docx")

        with open(input_pdf, "wb") as f:
            f.write(pdf_bytes)

        success = False
        translation_already_applied = False
        autofit = True
        all_errors = []

        # ── PRIMARY: PyMuPDF pdf_to_clean_docx.py (Clean inline text & borderless columns for Phrase) ──
        try:
            import sys as _sys
            _sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            from pdf_to_clean_docx import convert_pdf_to_clean_docx
            convert_pdf_to_clean_docx(input_pdf, clean_docx, tx_map if tx_map else None)
            if os.path.exists(clean_docx) and os.path.getsize(clean_docx) > 1000:
                success = True
                translation_already_applied = True
                print(f"pdf_to_clean_docx conversion succeeded: {os.path.getsize(clean_docx)} bytes")
        except Exception as e:
            all_errors.append(f"pdf_to_clean_docx: {e}")
            print(f"pdf_to_clean_docx failed: {e}")

        # ── FALLBACK 1: PyMuPDF layer-by-layer extraction (Standard Inline Flow for Phrase) ──
        if not success:
            try:
                import sys as _sys
                _sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                from pdf_layer_extractor import convert_pdf_to_docx as _layer_convert
                _layer_convert(input_pdf, clean_docx, tx_map if tx_map else None)
                if os.path.exists(clean_docx) and os.path.getsize(clean_docx) > 1000:
                    success = True
                    translation_already_applied = True
                    print(f"pdf_layer_extractor conversion succeeded: {os.path.getsize(clean_docx)} bytes")
            except Exception as e:
                all_errors.append(f"pdf_layer_extractor: {e}")
                print(f"pdf_layer_extractor failed: {e}")

        # ── FALLBACK 2: pdf_to_editable_docx (Exact-Replica background image + floating text boxes) ──
        if not success:
            try:
                import sys as _sys
                _sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                from pdf_to_editable_docx import convert_pdf_to_editable_docx
                convert_pdf_to_editable_docx(input_pdf, clean_docx, tx_map if tx_map else None)
                if os.path.exists(clean_docx) and os.path.getsize(clean_docx) > 1000:
                    success = True
                    translation_already_applied = True
                    print(f"pdf_to_editable_docx conversion succeeded: {os.path.getsize(clean_docx)} bytes")
            except Exception as e:
                all_errors.append(f"pdf_to_editable_docx: {e}")
                print(f"pdf_to_editable_docx failed: {e}")

        # ── FALLBACK 3: pdf2docx (standard editable conversion) ──
        if not success:
            try:
                cv = Converter(input_pdf)
                cv.convert(clean_docx)
                cv.close()
                if os.path.exists(clean_docx) and os.path.getsize(clean_docx) > 1000:
                    success = True
                    translation_already_applied = False
                    autofit = True
                    print(f"pdf2docx conversion succeeded: {os.path.getsize(clean_docx)} bytes")
            except Exception as e:
                all_errors.append(f"pdf2docx: {e}")
                print(f"pdf2docx failed: {e}")

        # ── FALLBACK 4: CloudConvert (OCR + conversion for scanned/image PDFs) ──
        if not success:
            cc_key = req.cloudConvertApiKey or os.environ.get("CLOUDCONVERT_API_KEY", "")
            try:
                requires_ocr = is_scanned_pdf(input_pdf)
                ocr_lang = get_cloudconvert_ocr_lang(req.sourceLanguage)
                export_url = await run_cloudconvert_job(
                    input_pdf, "pdf", "docx", cc_key,
                    ocr=requires_ocr, ocr_language=ocr_lang if requires_ocr else None
                )
                async with httpx.AsyncClient() as client:
                    download_r = await client.get(export_url)
                    with open(clean_docx, "wb") as f:
                        f.write(download_r.content)
                if os.path.exists(clean_docx) and os.path.getsize(clean_docx) > 1000:
                    success = True
                    translation_already_applied = False
                    autofit = False
                    print(f"CloudConvert conversion succeeded")
            except Exception as e:
                all_errors.append(f"CloudConvert: {e}")
                print(f"CloudConvert failed: {e}")

        if not success:
            raise HTTPException(
                status_code=500,
                detail=f"All PDF conversion methods failed: {'; '.join(all_errors)}"
            )

        # Apply translations if not already handled
        if tx_map and not translation_already_applied:
            apply_docx_translation(clean_docx, final_docx, tx_map, autofit=autofit)
        else:
            import shutil
            shutil.copy(clean_docx, final_docx)

        with open(final_docx, "rb") as f:
            output_bytes = f.read()

        return StreamingResponse(
            io.BytesIO(output_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={safe_filename}_translated.docx"}
        )


@app.get("/api/proxy-image")
async def proxy_image(url: str = Query(...)):
    if not url.startswith("http"):
        raise HTTPException(status_code=400, detail="Invalid URL.")
    async with httpx.AsyncClient() as client:
        r = await client.get(url)
        return Response(content=r.content, media_type=r.headers.get("content-type", "image/png"))

@app.post("/api/send-proposal")
async def send_proposal(req: SendProposalRequest):
    to_email = os.environ.get("PROPOSAL_DESTINATION_EMAIL", "")
    subject = f"📋 New translation proposal request: {req.filename}"
    
    # Simple SMTP email sending setup
    smtp_host = os.environ.get("SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(os.environ.get("SMTP_PORT", "465"))
    smtp_user = os.environ.get("SMTP_USER", "")
    smtp_pass = os.environ.get("SMTP_PASS", "")
    
    body_html = f"""
    <html>
    <body>
      <h2>Aura Translation Service - Proposal Request</h2>
      <p><strong>Name:</strong> {req.name}</p>
      <p><strong>Email:</strong> {req.email}</p>
      <p><strong>File Name:</strong> {req.filename}</p>
      <p><strong>Languages:</strong> {req.sourceLanguage} to {req.targetLanguage}</p>
      <h3>Message:</h3>
      <p>{req.message or 'No message provided'}</p>
    </body>
    </html>
    """

    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    msg = MIMEMultipart()
    msg['From'] = smtp_user or req.email
    msg['To'] = to_email
    msg['Subject'] = subject
    msg['Reply-To'] = req.email
    msg.attach(MIMEText(body_html, 'html'))

    # Attach SVG
    svg_att = MIMEBase('image', 'svg+xml')
    svg_att.set_payload(req.svgText.encode('utf-8'))
    encoders.encode_base64(svg_att)
    att_name = req.filename.replace(".eps", ".svg")
    if not att_name.endswith(".svg"):
        att_name += ".svg"
    svg_att.add_header('Content-Disposition', 'attachment', filename=att_name)
    msg.attach(svg_att)

    if smtp_user and smtp_pass:
        try:
            with smtplib.SMTP_SSL(smtp_host, smtp_port) as server:
                server.login(smtp_user, smtp_pass)
                server.sendmail(smtp_user or req.email, to_email, msg.as_string())
            return {"success": True, "mail_sent": True}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"SMTP Send failed: {e}")
    else:
        # Mock successful send when SMTP is not configured
        print("SMTP credentials not configured. Mocking success for proposal sending.")
        return {"success": True, "mail_sent": True, "mocked": True}

# --- Drop-in Router for api.php Compatibility ---

@app.api_route("/api.php", methods=["GET", "POST", "OPTIONS"])
async def api_php_proxy(request: Request, action: Optional[str] = Query(None)):
    if request.method == "OPTIONS":
        return Response()
        
    if not action:
        raise HTTPException(status_code=400, detail="Action parameter is required.")

    if action == "check-env":
        return await check_env()
        
    elif action == "log-login":
        body = await request.json()
        return await log_login(LogLoginRequest(**body))
        
    elif action == "ai-translate":
        body = await request.json()
        return await ai_translate(AiTranslateRequest(**body))
        
    elif action == "ocr-extract":
        body = await request.json()
        return await ocr_extract(OcrExtractRequest(**body))
        
    elif action == "pdf-to-docx":
        body = await request.json()
        return await pdf_to_docx(PdfToDocxRequest(**body))
        
    elif action == "proxy-image":
        url = request.query_params.get("url")
        return await proxy_image(url=url)
        
    elif action == "send-proposal":
        body = await request.json()
        return await send_proposal(SendProposalRequest(**body))
        
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported action: {action}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
